from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from zoneinfo import ZoneInfo
from difflib import SequenceMatcher
from io import BytesIO
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip().lower()


def _set_space_preserve(node):
    node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _first_run_properties(paragraph):
    for r in paragraph._p.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            return deepcopy(rPr)
    for change_tag in ("w:ins", "w:del"):
        for change in paragraph._p.findall(qn(change_tag)):
            for r in change.findall(qn("w:r")):
                rPr = r.find(qn("w:rPr"))
                if rPr is not None:
                    return deepcopy(rPr)
    return None


def _body_run_properties(paragraph):
    """Inherit font/size etc. but never force newly drafted body text to bold."""
    rPr = _first_run_properties(paragraph)
    if rPr is None:
        return None
    for tag in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag))
        if el is not None:
            rPr.remove(el)
    return rPr


def _word_timestamp():
    # Word commonly displays the literal revision clock from w:date.  Use Istanbul
    # wall-clock time without an offset to avoid the previous 3-hour UTC shift.
    return datetime.now(ZoneInfo("Europe/Istanbul")).replace(tzinfo=None).isoformat(timespec="seconds")


def _plain_run(text: str, rPr=None):
    r = OxmlElement("w:r")
    if rPr is not None:
        r.append(deepcopy(rPr))
    t = OxmlElement("w:t")
    _set_space_preserve(t)
    t.text = text
    r.append(t)
    return r


def _tracked_run(tag: str, text: str, change_id: int, author: str, rPr=None):
    change = OxmlElement(tag)
    change.set(qn("w:id"), str(change_id))
    change.set(qn("w:author"), author)
    change.set(qn("w:date"), _word_timestamp())
    r = OxmlElement("w:r")
    if rPr is not None:
        r.append(deepcopy(rPr))
    t_tag = "w:delText" if tag == "w:del" else "w:t"
    t = OxmlElement(t_tag)
    _set_space_preserve(t)
    t.text = text
    r.append(t)
    change.append(r)
    return change


def _clear_paragraph_content(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def _tokens(text: str):
    # Keep whitespace attached as tokens so Word redline remains readable.
    return re.findall(r"\s+|[\wÇĞİÖŞÜçğıöşü]+|[^\w\s]", text or "", flags=re.UNICODE)


def replace_paragraph_tracked(paragraph, new_text: str, change_id: int, author: str):
    """Minimal redline: preserve equal text, track only changed fragments."""
    old_text = paragraph.text
    rPr = _body_run_properties(paragraph)
    old_tokens, new_tokens = _tokens(old_text), _tokens(new_text)
    sm = SequenceMatcher(None, old_tokens, new_tokens, autojunk=False)
    _clear_paragraph_content(paragraph)
    p = paragraph._p

    for op, i1, i2, j1, j2 in sm.get_opcodes():
        old_piece = "".join(old_tokens[i1:i2])
        new_piece = "".join(new_tokens[j1:j2])
        if op == "equal":
            if new_piece:
                p.append(_plain_run(new_piece, rPr))
        elif op == "delete":
            if old_piece:
                p.append(_tracked_run("w:del", old_piece, change_id, author, rPr)); change_id += 1
        elif op == "insert":
            if new_piece:
                p.append(_tracked_run("w:ins", new_piece, change_id, author, rPr)); change_id += 1
        elif op == "replace":
            if old_piece:
                p.append(_tracked_run("w:del", old_piece, change_id, author, rPr)); change_id += 1
            if new_piece:
                p.append(_tracked_run("w:ins", new_piece, change_id, author, rPr)); change_id += 1
    return change_id


def insert_after_tracked(paragraph, new_text: str, change_id: int, author: str):
    new_p = OxmlElement("w:p")
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None:
        new_p.append(deepcopy(pPr))
    rPr = _body_run_properties(paragraph)
    new_p.append(_tracked_run("w:ins", new_text, change_id, author, rPr))
    paragraph._p.addnext(new_p)
    return change_id + 1


def append_end_tracked(doc, new_text: str, change_id: int, author: str):
    donor = next((p for p in reversed(doc.paragraphs) if p.text.strip()), None)
    p = doc.add_paragraph()
    if donor is not None:
        pPr = donor._p.find(qn("w:pPr"))
        if pPr is not None:
            p._p.insert(0, deepcopy(pPr))
        rPr = _body_run_properties(donor)
    else:
        rPr = None
    _clear_paragraph_content(p)
    p._p.append(_tracked_run("w:ins", new_text, change_id, author, rPr))
    return change_id + 1


def _all_paragraphs(doc):
    items = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                items.extend(cell.paragraphs)
    return items


def find_best_paragraph(doc, anchor_text: str):
    anchor = _norm(anchor_text)
    if not anchor:
        return None, 0.0
    paragraphs = _all_paragraphs(doc)
    for p in paragraphs:
        txt = _norm(p.text)
        if anchor in txt or (txt and txt in anchor and len(txt) > 25):
            return p, 1.0
    best_p, best_score = None, 0.0
    for p in paragraphs:
        txt = _norm(p.text)
        if not txt:
            continue
        score = SequenceMatcher(None, anchor, txt).ratio()
        if score > best_score:
            best_p, best_score = p, score
    return best_p, best_score


def highlight_placeholders(doc):
    patterns = [
        re.compile(r"[…]{3,}"), re.compile(r"\.{5,}"),
        re.compile(r"\[\s*(?:[A-ZÇĞİÖŞÜ0-9 _/-]{0,40})\s*\]"),
        re.compile(r"\bGG[./-]AA[./-]YYYY\b", re.I),
    ]
    count = 0
    for p in _all_paragraphs(doc):
        for run in list(p.runs):
            txt = run.text or ""
            spans = []
            for pat in patterns:
                spans.extend((m.start(), m.end()) for m in pat.finditer(txt))
            if not spans:
                continue
            spans.sort(); merged=[]
            for s,e in spans:
                if merged and s <= merged[-1][1]: merged[-1]=(merged[-1][0],max(merged[-1][1],e))
                else: merged.append((s,e))
            parent=run._r.getparent(); idx=parent.index(run._r); rPr=run._r.find(qn("w:rPr")); parent.remove(run._r)
            cursor=0; pieces=[]
            for s,e in merged:
                if s>cursor: pieces.append((txt[cursor:s],False))
                pieces.append((txt[s:e],True)); cursor=e
            if cursor<len(txt): pieces.append((txt[cursor:],False))
            for piece,hi in pieces:
                nr=OxmlElement("w:r")
                if rPr is not None: nr.append(deepcopy(rPr))
                if hi:
                    nrPr=nr.find(qn("w:rPr"))
                    if nrPr is None: nrPr=OxmlElement("w:rPr"); nr.insert(0,nrPr)
                    hl=OxmlElement("w:highlight"); hl.set(qn("w:val"),"yellow"); nrPr.append(hl); count+=1
                t=OxmlElement("w:t"); _set_space_preserve(t); t.text=piece; nr.append(t); parent.insert(idx,nr); idx+=1
    return count



def _highlight_paragraph(paragraph, color: str):
    """Apply Word highlight to the visible runs of a paragraph without changing font/size."""
    changed = 0
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        rPr = run._r.get_or_add_rPr()
        old = rPr.find(qn("w:highlight"))
        if old is not None:
            rPr.remove(old)
        hl = OxmlElement("w:highlight")
        hl.set(qn("w:val"), color)
        rPr.append(hl)
        changed += 1
    return changed


def highlight_flags(doc, flags):
    """
    ORANGE = risk/dikkat; BLUE = Rule Library ile anlamsal eşleşmeyen yeni hüküm.
    Uses paragraph highlighting only; no font/size/bold mutation.
    """
    stats = {"orange": 0, "blue": 0, "skipped": []}
    for flag in flags or []:
        color = (flag.get("color") or "").lower()
        if color not in ("orange", "blue"):
            continue
        anchor = (flag.get("anchor_text") or "").strip()
        p, score = find_best_paragraph(doc, anchor)
        if p is None or score < 0.42:
            stats["skipped"].append({**flag, "reason": f"Anchor bulunamadı ({score:.2f})"})
            continue
        # Word's supported highlight palette has no orange. darkYellow is used as the
        # closest native highlight and remains removable with Word's formatting controls.
        word_color = "darkYellow" if color == "orange" else "cyan"
        if _highlight_paragraph(p, word_color):
            stats[color] += 1
    return stats



def _semantic_change_ratio(old_text: str, new_text: str) -> float:
    old_tokens = re.findall(r"\w+|[^\w\s]", old_text or "", flags=re.UNICODE)
    new_tokens = re.findall(r"\w+|[^\w\s]", new_text or "", flags=re.UNICODE)
    if not old_tokens and not new_tokens:
        return 0.0
    sm = SequenceMatcher(None, old_tokens, new_tokens)
    return 1.0 - sm.ratio()


def choose_redline_mode(old_text: str, new_text: str, requested: str = "AUTO") -> str:
    """
    MICRO only when the edit is genuinely small.
    BLOCK when sentence structure materially changes, preventing word-by-word redline confetti.
    """
    req = (requested or "AUTO").upper()
    if req in {"MICRO", "BLOCK"}:
        return req
    ratio = _semantic_change_ratio(old_text, new_text)
    old_words = len((old_text or "").split())
    new_words = len((new_text or "").split())
    # Small edits: <= ~22% token change and broadly similar length.
    if ratio <= 0.22 and abs(old_words-new_words) <= max(4, int(max(old_words, new_words, 1)*0.18)):
        return "MICRO"
    return "BLOCK"


def replace_paragraph_block(doc, paragraph, new_text, author, date_str):
    """
    Delete the old paragraph text as one tracked block and insert the new paragraph text as one tracked block.
    Paragraph properties are preserved; inserted run copies the paragraph's first-run formatting.
    """
    p = paragraph._p
    old_text = paragraph.text
    # capture first visible run formatting before clearing
    template_rpr = None
    for rr in p.findall(qn("w:r")):
        rpr = rr.find(qn("w:rPr"))
        if rpr is not None:
            template_rpr = deepcopy(rpr)
            break

    # remove existing visible content but preserve pPr
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)

    del_el = OxmlElement("w:del")
    del_el.set(qn("w:author"), author)
    del_el.set(qn("w:date"), date_str)
    del_el.set(qn("w:id"), str(next_change_id(doc)))
    r_old = OxmlElement("w:r")
    if template_rpr is not None:
        r_old.append(deepcopy(template_rpr))
    dt = OxmlElement("w:delText")
    dt.set(qn("xml:space"), "preserve")
    dt.text = old_text
    r_old.append(dt)
    del_el.append(r_old)
    p.append(del_el)

    ins_el = OxmlElement("w:ins")
    ins_el.set(qn("w:author"), author)
    ins_el.set(qn("w:date"), date_str)
    ins_el.set(qn("w:id"), str(next_change_id(doc)))
    r_new = OxmlElement("w:r")
    if template_rpr is not None:
        r_new.append(deepcopy(template_rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = new_text
    r_new.append(t)
    ins_el.append(r_new)
    p.append(ins_el)


def apply_track_format_highlight(doc, paragraph, color: str, author: str, date_str: str):
    """
    Apply highlight AND record it as a tracked formatting change (w:rPrChange),
    so Word can show/accept/reject it as formatting rather than permanent paint.
    """
    changed = 0
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        rPr = run._r.get_or_add_rPr()
        # Snapshot prior formatting into rPrChange.
        previous = deepcopy(rPr)
        for old_change in previous.findall(qn("w:rPrChange")):
            previous.remove(old_change)
        # Current highlight.
        old_hl = rPr.find(qn("w:highlight"))
        if old_hl is not None:
            rPr.remove(old_hl)
        hl = OxmlElement("w:highlight")
        hl.set(qn("w:val"), color)
        rPr.append(hl)

        change = OxmlElement("w:rPrChange")
        change.set(qn("w:id"), str(next_change_id(doc)))
        change.set(qn("w:author"), author)
        change.set(qn("w:date"), date_str)
        change.append(previous)
        rPr.append(change)
        changed += 1
    return changed


def highlight_flags_tracked(doc, flags, author, date_str):
    stats = {"orange": 0, "blue": 0, "skipped": []}
    for flag in flags or []:
        color = (flag.get("color") or "").lower()
        if color not in ("orange", "blue"):
            continue
        anchor_text = (flag.get("anchor_text") or "").strip()
        p, score = find_best_paragraph(doc, anchor_text)
        if p is None or score < 0.42:
            stats["skipped"].append({**flag, "reason": f"Anchor bulunamadı ({score:.2f})"})
            continue
        word_color = "darkYellow" if color == "orange" else "cyan"
        if apply_track_format_highlight(doc, p, word_color, author, date_str):
            stats[color] += 1
    return stats

def apply_revisions_to_docx(original_bytes: bytes, revisions: list[dict], author: str = "Av. Onur Güneş", flags: list[dict] | None = None):
    doc=Document(BytesIO(original_bytes)); change_id=1000; applied=[]; skipped=[]
    for rev in revisions:
        action=(rev.get("action") or "REPLACE_PARAGRAPH").upper(); new_text=(rev.get("replacement_text") or "").strip(); anchor=(rev.get("anchor_text") or "").strip()
        if not new_text: skipped.append({**rev,"reason":"Boş revizyon metni"}); continue
        if action=="APPEND_END":
            change_id=append_end_tracked(doc,new_text,change_id,author); applied.append({**rev,"match_score":1.0}); continue
        paragraph,score=find_best_paragraph(doc,anchor)
        if paragraph is None or score<0.42: skipped.append({**rev,"reason":f"Anchor bulunamadı (eşleşme {score:.2f})"}); continue
        if action=="APPEND_AFTER": change_id=insert_after_tracked(paragraph,new_text,change_id,author)
        else: change_id=replace_paragraph_tracked(paragraph,new_text,change_id,author)
        applied.append({**rev,"match_score":score})
    placeholder_count=highlight_placeholders(doc)
    flag_stats=highlight_flags_tracked(doc, flags or [], author, date_str)
    bio=BytesIO(); doc.save(bio)
    return bio.getvalue(),applied,skipped,placeholder_count,flag_stats
