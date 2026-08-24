from __future__ import annotations

import json
import os
import re
from io import BytesIO
from pathlib import Path

from docx import Document
from openai import OpenAI
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parent


def get_api_key() -> str | None:
    try:
        import streamlit as st
        if "OPENAI_API_KEY" in st.secrets:
            return st.secrets["OPENAI_API_KEY"]
    except Exception:
        pass
    return os.getenv("OPENAI_API_KEY")


def get_model() -> str:
    try:
        import streamlit as st
        return st.secrets.get("OPENAI_MODEL", "gpt-5-mini")
    except Exception:
        return os.getenv("OPENAI_MODEL", "gpt-5-mini")


def extract_text(filename: str, raw: bytes) -> str:
    lower=filename.lower()
    if lower.endswith(".docx"):
        doc=Document(BytesIO(raw))
        parts=[]
        for p in doc.paragraphs:
            t=p.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells=[" ".join(c.text.split()) for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if lower.endswith(".pdf"):
        reader=PdfReader(BytesIO(raw))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    raise ValueError("Yalnız DOCX ve PDF desteklenir.")


def load_actor_rules():
    return json.loads((ROOT/"oli_actor_tv_rules.json").read_text(encoding="utf-8"))


def load_revision_dna():
    path=ROOT/"oli_revision_dna.json"
    return json.loads(path.read_text(encoding="utf-8")) if path.exists() else {}

def clean_json(text: str):
    """Parse model JSON robustly; tolerate fences/prose and repair once if needed."""
    raw=(text or "").strip()
    raw=re.sub(r"^```(?:json)?\s*","",raw,flags=re.I)
    raw=re.sub(r"\s*```$","",raw)
    first=raw.find("{")
    last=raw.rfind("}")
    if first>=0 and last>first:
        raw=raw[first:last+1]
    raw=raw.replace("\u00a0"," ")
    raw=re.sub(r",\s*([}\]])", r"\1", raw)

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        api_key=get_api_key()
        if not api_key:
            raise
        client=OpenAI(api_key=api_key)
        repair=client.responses.create(
            model=get_model(),
            input=[
                {
                    "role":"system",
                    "content":(
                        "Aşağıdaki metin hukuki analiz içeren bozuk JSON'dur. "
                        "Yalnız JSON sözdizimini düzelt. Hiçbir hukuki içeriği, kelimeyi, "
                        "madde numarasını veya değeri değiştirme/ekleme/silme. "
                        "Yanıt yalnız geçerli JSON olsun."
                    )
                },
                {"role":"user","content":raw}
            ],
        )
        repaired=(repair.output_text or "").strip()
        repaired=re.sub(r"^```(?:json)?\s*","",repaired,flags=re.I)
        repaired=re.sub(r"\s*```$","",repaired)
        f=repaired.find("{")
        l=repaired.rfind("}")
        if f>=0 and l>f:
            repaired=repaired[f:l+1]
        repaired=re.sub(r",\s*([}\]])", r"\1", repaired)
        return json.loads(repaired)


def _basic_focus(contract_type: str) -> str:
    if contract_type=="Senarist Sözleşmesi":
        return """Teslim ve revizyon yükümlülükleri, opsiyon, ücret/hak kazanma, FSEK mali hak devri,
        jenerik/kredi, yeniden yazım, üçüncü kişi müdahaleleri, süre, fesih, yapımcının projeyi gerçekleştirmemesi,
        başka projelerde çalışma ve ek kullanım alanlarına özellikle bak."""
    if contract_type=="Yönetmen Sözleşmesi":
        return """Hazırlık/çekim takvimi, yaratıcı yetki, kurgu ve final cut süreci, yeniden çekim,
        tanıtım/gala, ücret ve ödeme, FSEK bağlantılı haklar, jenerik, seyahat/konaklama,
        fesih, proje ertelenmesi/askıya alınması ve başka projelerde çalışma alanlarına özellikle bak."""
    return """Münhasırlık, başka projede çalışma, çekim/set koşulları, ücret/ödeme, yeni sezon,
    cezai şart, fesih, reklam/ürün yerleştirme, tanıtım, sosyal medya, fiziksel görünüm,
    mali haklar, meslek birliği telifleri ve yayınlanmayan bölüm ücretlerine özellikle bak."""


def analyze_contract(
    contract_text: str,
    contract_type: str,
    project_type: str,
    negotiation_power: str,
    initial_note: str = "",
):
    api_key=get_api_key()
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY tanımlı değil.")

    revision_dna=load_revision_dna()

    rules=[]
    profile_note=""
    if contract_type=="Oyuncu Sözleşmesi" and project_type=="Ana Akım TV":
        rules=load_actor_rules()
        profile_note="Bu profil için aşağıdaki 30 kontrollü Opus kuralını esas al."
    else:
        profile_note=(
            "Bu kombinasyon için henüz özel Rule Library tamamlanmadı. "
            "Genel entertainment/media sözleşme incelemesi yap; özel kullanıcı standardı uydurma."
        )

    rules_text="\n".join(
        f"{r['id']} | {r['title']} | {r['standard']} | Öncelik:{r.get('priority','')}"
        for r in rules
    ) or "Özel kural seti yok."

    system=f"""Sen OLI Clean sözleşme inceleme motorusun.
Kullanıcı bir entertainment/media avukatıdır. Görevin sözleşmeyi yeniden yazmak değil;
müdahale edilmesi gereken yerleri doğru tespit edip kullanılabilir revizyon önerileri vermektir.

SÖZLEŞME TÜRÜ: {contract_type}
PROJE TÜRÜ: {project_type}
PAZARLIK GÜCÜ: {negotiation_power}
DOSYAYA ÖZGÜ İLK NOT / AJANS NOTU: {initial_note or "Yok"}

{profile_note}

İNCELEME ODAĞI:
{_basic_focus(contract_type)}

KONTROLLÜ KURALLAR:
{rules_text}

REVİZYON DNA — AV. ONUR GÜNEŞ ÇALIŞMA BİÇİMİ:
{json.dumps(revision_dna, ensure_ascii=False)}

REVİZYON KARAR TESTİ:
Her suggested_revision üretmeden önce içsel olarak şunu kontrol et: "Bu hukuki sonucu mevcut cümlenin yapısını koruyarak daha az kelime değiştirerek elde edebilir miyim?" Cevap evetse mikro müdahaleyi kullan. Tam hüküm yeniden yazımı son çaredir. Koruma hiç yoksa EK HÜKÜM önerebilirsin. DNA örneklerindeki rakamları/dosyaya özgü ayrıntıları otomatik kopyalama.

ZORUNLU İLKELER:
1. Sözleşmede olmayan bir madde numarası veya olgu uydurma.
2. Her tespit için sözleşmedeki gerçek madde/referansı belirt.
3. Gereksiz stil düzeltmelerini risk/revizyon olarak çıkarma.
4. Önerilen revizyon, mümkün olduğunca gelen sözleşmenin mevcut cümle yapısını korusun.
5. Birkaç kelimeyle hukuki sonuç düzeltilebiliyorsa yepyeni paragraf yazma.
6. Yeni koruyucu hüküm gerçekten eksikse bunu "EK HÜKÜM" olarak açıkça söyle.
7. Dosyaya özgü İlk Not genel standarttan daha öncelikli olabilir; fakat notta yazmayan bilgiyi uydurma.
8. Kullanıcı Word'e revizyonu kendisi işleyecek. Öneri kısa, net, kopyalanabilir ve sözleşme diline uyumlu olsun.
9. Aynı hukuki sorunu tekrar eden birden fazla bulguya bölme.
10. Yalnız anlamlı müdahaleleri getir.
11. WORD UYUMLULUK KURALI:
    - Mevcut bir sözleşme maddesi revize ediliyorsa `reference` alanına sözleşmede gerçekten geçen TEK madde numarasını yaz (örn. "4.11"). "4.11/4.12", "Madde 4", "ilgili madde", başlık veya açıklama yazma.
    - `current_excerpt` alanı sözleşmeden BİREBİR ve ARDIŞIK alınmış kısa bir parçadır; özetleme, sadeleştirme veya yeniden yazma. Mümkünse 8-35 kelime arasında olsun.
    - Mevcut hüküm üzerinde değişiklik yapılıyorsa `revision_type` yalnız MİKRO, CÜMLECİK, TAM MADDE veya SİLME olabilir. `EK HÜKÜM` yalnız sözleşmede gerçekten hiç bulunmayan yeni bir koruma için kullanılabilir.
    - `suggested_revision`, mevcut hükmün yerine Word'e işlenebilecek gerçek revizyon metni olmalıdır; "daraltılmalı", "eklenmeli", "önerilir" gibi talimat cümlesi yazma.
    - MİKRO/CÜMLECİK tercihinde mümkünse current_excerpt içindeki cümle yapısını koru ve yalnız gereken kelime/ibareleri değiştir.
    - Aynı mevcut maddeyi yanlışlıkla EK HÜKÜM olarak sınıflandırma.
12. short_notes alanı ajansa ilk gönderilecek kısa pazarlık notudur. Her not 2-3 kısa cümle olsun ve mutlaka şu üç unsuru içersin:
    (a) mevcut maddede ne düzenlendiği,
    (b) bunun neden sorunlu / fazla geniş / eksik olduğu,
    (c) nasıl değiştirilmesini istediğimiz.
    Hukuki mütalaa uzunluğuna çıkma; tam revizyon metnini burada tekrar etme.
    Örnek ton: "Mevcut düzenlemede OYUNCU'nun başka televizyon projelerinde yer alması tamamen yasaklanmış. Bu yasak fazla geniş. Yasağın yalnızca ana projenin çekimlerini aksatacak veya doğrudan rekabet yaratacak işler bakımından sınırlandırılmasını öneriyoruz."

YANIT YALNIZ GEÇERLİ JSON:
ÇIKTI KURALLARI:
- Markdown code fence kullanma.
- JSON dışında tek kelime yazma.
- Metin değerlerinin içindeki çift tırnakları \" ile escape et.
- Metin değerlerinde gerçek satır sonu kullanma; gerekiyorsa \n kullan.
- Trailing comma kullanma.

{{
  "overall_risk":"Düşük|Orta|Yüksek|Çok Yüksek",
  "executive_summary":"en fazla 4 cümle",
  "short_notes":[
    {{"reference":"5.4","note":"Ajansa gönderilebilir tek cümlelik kısa not"}}
  ],
  "revisions":[
    {{
      "reference":"5.4",
      "title":"Münhasırlık",
      "severity":"KRİTİK|ÖNEMLİ|DİKKAT",
      "problem":"Neden müdahale edilmeli?",
      "current_excerpt":"Sözleşmeden birebir, ardışık ve kısa alıntı; paraphrase etme",
      "revision_type":"MİKRO|CÜMLECİK|TAM MADDE|EK HÜKÜM|SİLME",
      "suggested_revision":"Word'e doğrudan işlenebilir revizyon metni; talimat değil gerçek metin"
    }}
  ]
}}
"""

    client=OpenAI(api_key=api_key)
    resp=client.responses.create(
        model=get_model(),
        input=[
            {"role":"system","content":system},
            {"role":"user","content":"SÖZLEŞME METNİ:\n"+contract_text[:150000]},
        ],
    )
    return clean_json(resp.output_text)
