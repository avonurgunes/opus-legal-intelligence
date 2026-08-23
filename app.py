import streamlit as st
import json
import re
from datetime import datetime
from zoneinfo import ZoneInfo
from difflib import SequenceMatcher
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from openai import OpenAI

from revision_engine import apply_revisions_to_docx
from learning_engine import build_candidates, approve_candidates, load_memory, learning_prompt_block, match_batch_documents, batch_candidates, cluster_candidates
from mediation import render_mediation

st.set_page_config(
    page_title="Opus Legal Intelligence",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

RULES = json.loads(Path(__file__).with_name("rules.json").read_text(encoding="utf-8"))
REVISION_LIBRARY = json.loads(
    Path(__file__).with_name("revision_library.json").read_text(encoding="utf-8")
)
CLAUSE_BANK = json.loads(
    Path(__file__).with_name("clause_bank.json").read_text(encoding="utf-8")
)

POWER_GUIDANCE = {
    "Düşük": "Sadece kritik konuları masaya taşı; ikincil risklerde pazarlık sermayesini koru.",
    "Orta": "Kritik riskleri ve önemli ticari risklerin çoğunu müzakere et.",
    "Yüksek": "Opus standartlarına geniş ölçüde yaklaş; yüksek ve orta risklerde güçlü revizyon iste.",
    "Çok Yüksek": "İdeal Opus pozisyonuna mümkün olduğunca yaklaş; gereksiz tek taraflı hükümleri kabul etme.",
}

REV_MODE_BY_POWER = {
    "Düşük": "MINIMUM",
    "Orta": "STANDARD",
    "Yüksek": "STANDARD",
    "Çok Yüksek": "STRONG",
}

CSS = """
<style>
:root{
  --opus-black:#111214;
  --opus-panel:#1a1b1f;
  --opus-gold:#b89550;
  --opus-gold2:#d2b36f;
  --opus-cream:#f4efe5;
}
.stApp{
  background:
    radial-gradient(circle at 10% 0%, rgba(184,149,80,.10), transparent 28%),
    linear-gradient(180deg,#0f1012 0%,#15161a 100%);
}
[data-testid="stHeader"]{background:rgba(15,16,18,.75);}
.block-container{max-width:1420px;padding-top:2rem;}
h1,h2,h3{letter-spacing:-.02em}
h1{color:var(--opus-cream)!important;font-weight:800!important}
h2,h3{color:#f3f1ec!important}
p, label, .stMarkdown, [data-testid="stCaptionContainer"]{color:#d9d5cc!important}
.opus-kicker{color:var(--opus-gold2);letter-spacing:.16em;font-size:.78rem;font-weight:700}
.opus-hero{
  padding:1.7rem 1.9rem;border:1px solid rgba(184,149,80,.35);
  background:linear-gradient(135deg,rgba(184,149,80,.13),rgba(255,255,255,.025));
  border-radius:18px;margin-bottom:1.2rem
}
.opus-module{
  min-height:125px;padding:1.1rem;border-radius:16px;
  border:1px solid rgba(184,149,80,.20);background:rgba(255,255,255,.035)
}
.opus-module.active{border-color:rgba(184,149,80,.65);background:rgba(184,149,80,.09)}
.opus-module-title{font-size:1.12rem;color:#fff;font-weight:750;margin-bottom:.4rem}
.opus-module-sub{font-size:.87rem;color:#aaa59b}
div[data-testid="stMetric"]{
  border:1px solid rgba(184,149,80,.22);border-radius:14px;padding:.8rem;
  background:rgba(255,255,255,.025)
}
div.stButton > button, div.stDownloadButton > button{
  border-radius:10px;border:1px solid rgba(184,149,80,.6);
}
div.stButton > button[kind="primary"]{
  background:linear-gradient(90deg,#a98543,#c5a563);color:#0f1012;border:none;font-weight:800
}
[data-testid="stFileUploader"]{border-radius:14px}
details{border:1px solid rgba(184,149,80,.15)!important;border-radius:12px!important}
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)


def read_docx(file_bytes: bytes) -> str:
    doc = Document(BytesIO(file_bytes))
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    parts = []
    for i, page in enumerate(reader.pages, 1):
        txt = page.extract_text() or ""
        if txt.strip():
            parts.append(f"\n--- SAYFA {i} ---\n{txt}")
    return "\n".join(parts)


def extract_text(uploaded) -> str:
    data = uploaded.getvalue()
    name = uploaded.name.lower()
    if name.endswith(".docx"):
        return read_docx(data)
    if name.endswith(".pdf"):
        return read_pdf(data)
    raise ValueError("Desteklenmeyen dosya türü.")


def get_api_key():
    try:
        return st.secrets["OPENAI_API_KEY"]
    except Exception:
        return None


def get_model():
    default = "gpt-5.6-terra"
    try:
        return st.secrets.get("OPENAI_MODEL", default)
    except Exception:
        return default


def clean_json(text: str):
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.I)
    text = re.sub(r"\s*```$", "", text)
    start = text.find("{")
    end = text.rfind("}")
    if start >= 0 and end > start:
        text = text[start:end+1]
    return json.loads(text)


def analyse_contract(contract_text: str, negotiation_power: str, initial_note: str = ""):
    rules_text = "\n".join(
        f"{r['id']} | {r['title']} | Öncelik: {r['priority']} | Opus standardı: {r['standard']}"
        for r in RULES
    )

    system = """Sen Opus Legal Intelligence (OLI) sözleşme analiz motorusun.
Ana akım TV dizisi oyuncu sözleşmesini yalnız verilen Opus kural setine göre incele.

İLKELER:
- Müvekkil OYUNCU/AJANS.
- Avantajlı hüküm GREEN/KORU olabilir.
- Hüküm yoksa NOT_FOUND; fakat kural koruyucu bir hükmün sözleşmede bulunmasını gerektiriyorsa assessment içinde bunun eksiklik olduğunu açıkla.
- 'Geçmişte revize edilmemiş = kabul edilmiş' varsayımı yapma.
- Pazarlık gücü hukuki riski değil müzakere önceliğini etkiler.
- Metinde desteklenmeyen bilgi üretme.
- Yanıt yalnız geçerli JSON.

ŞEMA:
{
 "overall_risk":"LOW|MEDIUM|HIGH|VERY_HIGH",
 "executive_summary":"kısa özet",
 "top_negotiation_points":["en fazla 5"],
 "findings":[{
   "rule_id":"OLI-TV-001",
   "title":"...",
   "status":"RED|ORANGE|YELLOW|GREEN|NOT_FOUND",
   "contract_reference":"madde no/bölüm",
   "clause_excerpt":"kısa hüküm özeti",
   "assessment":"neden",
   "recommended_revision":"kısa prensip",
   "negotiation_priority":"MUST|SHOULD|OPTIONAL|KEEP",
   "confidence":"HIGH|MEDIUM|LOW"
 }]
}
Her 30 kural için findings üret.
"""

    user = f"""PROFİL: ACTOR_TV_MAINSTREAM
PAZARLIK GÜCÜ: {negotiation_power}
STRATEJİ: {POWER_GUIDANCE[negotiation_power]}
DOSYAYA ÖZGÜ İLK NOT / AJANS NOTU: {initial_note or "Yok"}
Bu not yalnız bu dosyanın analiz ve müzakere önceliklerini etkiler; Opus standardı değildir.

ONAYLANMIŞ GEÇMİŞ DRAFTING ÖĞRENİMLERİ:
{learning_prompt_block(contract_text)}
Bunlar kalıp cümle değildir. Gelen sözleşmenin terminolojisini ve cümle yapısını koruyarak yalnız davranış/tercih örneği olarak kullan.

RULE LIBRARY:
{rules_text}

SÖZLEŞME:
{contract_text[:140000]}
"""

    client = OpenAI(api_key=get_api_key())
    resp = client.responses.create(
        model=get_model(),
        input=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
    )
    return clean_json(resp.output_text)


def revision_context_for(rule_id: str):
    for entry in REVISION_LIBRARY.get("entries", []):
        if entry.get("rule_id") == rule_id:
            return entry
    return None


def clause_bank_for(rule_id: str):
    return [e for e in CLAUSE_BANK.get("entries", []) if e.get("rule_id") == rule_id]


def build_revision_drafts(contract_text, selected_findings, negotiation_power):
    lib_parts = []
    bank_parts = []
    for f in selected_findings:
        entry = revision_context_for(f.get("rule_id"))
        if entry:
            lib_parts.append(json.dumps(entry, ensure_ascii=False))
        for bank_entry in clause_bank_for(f.get("rule_id")):
            bank_parts.append(json.dumps(bank_entry, ensure_ascii=False))

    system = """Sen OLI Revision Engine'sin.
Görevin sözleşmedeki seçilmiş bulgular için OPUS tarzında uygulanabilir revizyon metni üretmektir.

ÖNCELİK:
1) Varsa MADDE BANKASI'ndaki Av. Onur Güneş tarafından gözden geçirilmiş hazır cümle.
2) Varsa REVISION LIBRARY'deki doğrulanmış Opus yaklaşımı.
3) Rule Library'deki standardın özü.
4) Mevcut sözleşmenin terminolojisi, taraf tanımları, proje adı ve madde dili.

KURALLAR:
- Yeni hukuki politika icat etme.
- MADDE BANKASI'nda uygun cümle varsa esas metin olarak onu kullan; yalnız sözleşmenin tanımlarına, proje adına, madde numarasına ve mevcut cümlenin gramerine uyacak kadar değiştir.
- Madde Bankası metnini gereksiz yere uzatma, yeni şartlar ekleme veya daha ayrıntılı hale getirme.
- Mevcut hüküm kısmen uygunsa paragrafı baştan yazmak yerine mümkün olan en küçük kelime/cümlecik değişikliğiyle Madde Bankası standardını mevcut cümleye yedir.
- Geçmiş Opus metnini körü körüne kopyalama; mevcut sözleşmeye uyarla.
- Mevcut hüküm varsa action=REPLACE_PARAGRAPH ve anchor_text sözleşmeden birebir kısa bir parça olsun.
- Koruyucu hüküm tamamen eksikse NOT_FOUND diye bırakma: mutlaka yeni hüküm üret. Sözleşmede konu bakımından en uygun mevcut maddeyi anchor seç ve APPEND_AFTER kullan. APPEND_END yalnız gerçekten uygun bölüm bulunamıyorsa son çaredir.
- Eksik hükmün ekleneceği yeri konu bütünlüğüne göre seç: ücret/ödeme hükümleri ücret bölümüne; mali hak/telif hükümleri mali haklar bölümüne; fesih/cezai şart hükümleri ilgili fesih/ceza bölümüne; vergi/damga vergisi diğer hükümler/vergi bölümüne.
- APPEND_AFTER kullanırken anchor_text mutlaka sözleşmede BİREBİR bulunan ve hedef bölümdeki son uygun paragraftan 25-100 karakterlik bir parça olmalı. Madde numarasını uydurma. Yeni metnin numaralandırması belirsizse numarasız koruyucu paragraf üret; Word'e yerleştirildikten sonra kullanıcı kontrol eder.
- replacement_text sadece sözleşmeye girecek nihai madde/paragraf metni olsun.
- Pazarlık gücü sadece sertlik düzeyini belirler.
- Yanıt yalnız geçerli JSON.

ŞEMA:
{
 "revisions":[{
   "rule_id":"...",
   "title":"...",
   "action":"REPLACE_PARAGRAPH|APPEND_AFTER|APPEND_END",
   "anchor_text":"sözleşmeden birebir 15-120 karakter veya boş",
   "replacement_text":"Word'e işlenecek revize hüküm",
   "reason":"kısa açıklama",
   "confidence":"HIGH|MEDIUM|LOW"
 }]
}
"""

    mode = REV_MODE_BY_POWER[negotiation_power]
    user = f"""PAZARLIK GÜCÜ: {negotiation_power}
REVİZYON MODU: {mode}

SEÇİLMİŞ BULGULAR:
{json.dumps(selected_findings, ensure_ascii=False, indent=2)}

MADDE BANKASI:
{chr(10).join(bank_parts) if bank_parts else "Bu kurallar için onaylı hazır madde yok."}

REVISION LIBRARY:
{chr(10).join(lib_parts) if lib_parts else "Bu kurallar için özel geçmiş kalıp yok."}

SÖZLEŞME:
{contract_text[:140000]}
"""

    client = OpenAI(api_key=get_api_key())
    resp = client.responses.create(
        model=get_model(),
        input=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
    )
    return clean_json(resp.output_text)


def icon(status):
    return {
        "RED": "🔴", "ORANGE": "🟠", "YELLOW": "🟡",
        "GREEN": "🟢", "NOT_FOUND": "⚪"
    }.get(status, "⚪")



def analyse_extra_risks(contract_text: str):
    system = """Sen OLI Extra Risk Scanner'sın.
30 maddelik Opus Rule Library DIŞINDA kalan sözleşmesel riskleri ara.
Rule Library'deki konuları tekrar etme. Yalnız gerçekten yeni ve anlamlı hukuki/ticari bulguları yaz.
Bulguları otomatik revize etme; kullanıcıya not olarak sun.
Yanıt yalnız JSON:
{"extra_findings":[{"title":"...","risk":"HIGH|MEDIUM|LOW","reference":"madde/bölüm","assessment":"...","suggested_action":"..."}]}
Hiç ek bulgu yoksa boş liste döndür."""
    user = f"""MEVCUT 30 KURAL BAŞLIĞI:
{chr(10).join(r.get('title','') for r in RULES)}

SÖZLEŞME:
{contract_text[:140000]}"""
    client = OpenAI(api_key=get_api_key())
    resp = client.responses.create(
        model=get_model(),
        input=[{"role":"system","content":system},{"role":"user","content":user}]
    )
    return clean_json(resp.output_text)


def compare_final_revision(oli_text: str, final_text: str):
    system = """Sen OLI Learning Review motorusun.
OLI'nin ürettiği revize metin ile Av. Onur Güneş'in nihai revize metnini karşılaştır.
Yalnız anlamlı hukuki/drafting farklarını çıkar.
Bir hükme dokunulmamasını asla otomatik olarak 'Opus standardı' kabul etme.
Her fark için bunun kütüphaneye alınabilecek genel bir tercih mi, yoksa dosyaya özgü mü olabileceğini öner.
Yanıt yalnız JSON:
{"learning_candidates":[{"title":"...","oli_position":"...","final_position":"...","difference":"...","recommendation":"LIBRARY_CANDIDATE|CASE_SPECIFIC","confidence":"HIGH|MEDIUM|LOW"}]}"""
    user = f"""OLI REVİZE METİN:
{oli_text[:100000]}

AV. ONUR GÜNEŞ NİHAİ METİN:
{final_text[:100000]}"""
    client = OpenAI(api_key=get_api_key())
    resp = client.responses.create(
        model=get_model(),
        input=[{"role":"system","content":system},{"role":"user","content":user}]
    )
    return clean_json(resp.output_text)



def build_initial_review_note(result):
    system = """Sen OLI Ajans Bilgilendirme Notu motorusun.
Verilen sözleşme analizinden yalnız müdahale edilmesi, ajansa bildirilmesi veya karar alınması gereken noktaları kısa Türkçe not halinde yaz.
Her satır mümkünse sözleşme madde numarasıyla başlasın.
Dil kısa ve pratik olsun: '5.4. Münhasırlık Yapımcı onayına bağlanmış. Bilgilendirmeye çevirelim.' gibi.
Uzun hukuki açıklama yapma. GREEN/KORU maddelerini yazma. Yanıt yalnız JSON:
{"note_items":[{"reference":"5.4","title":"Münhasırlık","note":"..."}]}"""
    client = OpenAI(api_key=get_api_key())
    resp = client.responses.create(
        model=get_model(),
        input=[{"role":"system","content":system},
               {"role":"user","content":json.dumps(result, ensure_ascii=False)[:100000]}]
    )
    return clean_json(resp.output_text)


def classify_word_flags(contract_text: str, result: dict, extra_risks: dict | None = None):
    rules_text = "\n".join(f"{r['id']} | {r['title']} | {r['standard']}" for r in RULES)
    system = """Sen OLI Word Flag motorusun.
Sözleşmedeki yalnız iki tür hükmü işaretle:
ORANGE: hukuki/ticari risk veya manuel dikkat gerektiren hüküm; fakat Word revizyonu ile zaten doğrudan değiştirilecek parça olmak zorunda değil.
BLUE: mevcut Rule Library ile ANLAMSAL olarak eşleşmeyen, gerçekten yeni/öğrenilmemiş bir hukuki düzenleme.
Aynı konunun farklı kelimelerle yazılması BLUE değildir. Örneğin farklı yazılmış cezai şart yine cezai şarttır.
Her flag için anchor_text sözleşmede BİREBİR geçen 25-100 karakterlik kısa parça olmalı.
Yanıt yalnız JSON:
{"flags":[{"color":"orange|blue","title":"...","reference":"...","anchor_text":"...","reason":"..."}]}"""
    user = f"""RULE LIBRARY:
{rules_text}

30 KURAL ANALİZİ:
{json.dumps(result, ensure_ascii=False)}

EK RİSKLER:
{json.dumps(extra_risks or {}, ensure_ascii=False)}

SÖZLEŞME:
{contract_text[:140000]}"""
    client = OpenAI(api_key=get_api_key())
    resp = client.responses.create(model=get_model(), input=[{"role":"system","content":system},{"role":"user","content":user}])
    return clean_json(resp.output_text)


def compare_counterparty_return(our_text: str, returned_text: str):
    system = """Sen OLI Negotiation Compare motorusun.
Av. Onur Güneş tarafından gönderilen revize sözleşme ile karşı taraftan dönen sözleşmeyi hukuki anlam bakımından karşılaştır.
Sınıflar:
ACCEPTED = bizim değişiklik kabul edilmiş.
PARTIAL = kısmen kabul edilmiş/değiştirilmiş.
REJECTED = bizim değişiklik kaldırılmış/eski pozisyona dönülmüş.
NEW = karşı taraf yeni bir hüküm eklemiş.
Kısa, ajansa aktarılabilir dil kullan. Yanıt yalnız JSON:
{"items":[{"reference":"...","title":"...","status":"ACCEPTED|PARTIAL|REJECTED|NEW","our_position":"...","counterparty_position":"...","note":"..."}]}"""
    client = OpenAI(api_key=get_api_key())
    resp = client.responses.create(
        model=get_model(),
        input=[{"role":"system","content":system},
               {"role":"user","content":f"BİZİM GÖNDERDİĞİMİZ:\\n{our_text[:100000]}\\n\\nKARŞI TARAF DÖNÜŞÜ:\\n{returned_text[:100000]}"}]
    )
    return clean_json(resp.output_text)




def _norm_ref(ref):
    return (ref or "").strip().rstrip(".").replace("MADDE ","").replace("Madde ","").strip()

def extract_full_original_clause(contract_text, item):
    text=contract_text or ""; ref=_norm_ref(item.get("reference") or item.get("rule_id") or "")
    lines=text.splitlines()
    if ref and re.fullmatch(r"\d+(?:\.\d+)*",ref):
        pat=re.compile(rf"^\s*(?:MADDE\s+)?{re.escape(ref)}(?:[\.\-\):]|\s)",re.I)
        start=next((i for i,x in enumerate(lines) if pat.search(x)),None)
        if start is not None:
            depth=ref.count("."); collected=[lines[start]]
            nxt=re.compile(r"^\s*(?:MADDE\s+)?(\d+(?:\.\d+)*)(?:[\.\-\):]|\s)",re.I)
            for line in lines[start+1:]:
                m=nxt.search(line)
                if m and m.group(1)!=ref and m.group(1).count(".")<=depth: break
                collected.append(line)
            full="\n".join(collected).strip()
            if full: return full
    target=(item.get("target_text") or item.get("anchor_text") or "").strip()
    if target:
        for line in lines:
            if target in line or (len(target)>35 and target[:35] in line): return line.strip()
    return target or "Orijinal madde bulunamadı."

def build_full_revised_clause(original,item):
    if (item.get("mode") or "REPLACE").upper()=="APPEND_AFTER":
        return (original.rstrip()+"\n\n"+(item.get("append_text") or "").strip()).strip()
    repl=(item.get("replacement_text") or "").strip(); target=(item.get("target_text") or "").strip()
    if target and target in original: return original.replace(target,repl,1)
    return repl or original

def _preview_default_mode(item):
    req = (item.get("redline_mode") or "AUTO").upper()
    if item.get("mode") == "APPEND_AFTER":
        return "➕ Yeni hüküm"
    if req == "MICRO":
        return "🩹 Mikro değişiklik"
    if req == "PHRASE":
        return "✂️ Cümlecik / doğal ekleme"
    if req == "BLOCK":
        return "📝 Tam blok değişikliği"
    return "🤖 Otomatik seç"


def render_revision_preview(items, contract_text):
    st.subheader("Revizyon Önizleme")
    st.caption("Gerekçe → tam orijinal madde → tam revize madde. Word yalnız kabul ettiğin son metinlerden oluşturulur.")
    edited=[]; accepted=rejected=manual=0
    mm={"🤖 Otomatik seç":"AUTO","🩹 Mikro değişiklik":"MICRO","✂️ Cümlecik / doğal ekleme":"PHRASE","📝 Tam blok değişikliği":"BLOCK","➕ Yeni hüküm":"AUTO"}
    for i,item in enumerate(items):
        rid=item.get("reference") or item.get("rule_id") or f"REV-{i+1}"
        title=item.get("title") or item.get("rule_title") or rid
        original=extract_full_original_clause(contract_text,item)
        proposed=build_full_revised_clause(original,item)
        with st.expander(f"{i+1}. {rid} — {title}",expanded=(i<3)):
            st.markdown("**Neden revize ediyorum?**")
            st.info(item.get("reason") or item.get("problem") or "Revizyon gerekçesi belirtilmemiş.")
            st.markdown("**ORİJİNAL MADDE**")
            st.text_area("Orijinal",original,height=150,disabled=True,key=f"orig_{i}",label_visibility="collapsed")
            st.markdown("**REVİZE MADDE**")
            final=st.text_area("Revize",proposed,height=180,key=f"revised_{i}",label_visibility="collapsed")
            c1,c2=st.columns(2)
            decision=c1.radio("Karar",["✅ Kabul","❌ Reddet"],horizontal=True,key=f"decision_{i}")
            choices=["🤖 Otomatik seç","🩹 Mikro değişiklik","✂️ Cümlecik / doğal ekleme","📝 Tam blok değişikliği"]
            if item.get("mode")=="APPEND_AFTER": choices=["➕ Yeni hüküm","🤖 Otomatik seç"]
            default=_preview_default_mode(item); idx=choices.index(default) if default in choices else 0
            vm=c2.selectbox("Word'de uygulanma biçimi",choices,index=idx,key=f"vmode_{i}")
            if decision=="❌ Reddet": rejected+=1; continue
            accepted+=1; ni=dict(item); ni["redline_mode"]=mm[vm]
            ni["preview_original_clause"]=original; ni["preview_final_clause"]=final
            if (item.get("mode") or "REPLACE").upper()=="APPEND_AFTER":
                ni["append_text"]=final[len(original):].strip() if final.startswith(original) else final
            else:
                ni["target_text"]=original; ni["replacement_text"]=final
            if final.strip()!=proposed.strip(): manual+=1; ni["edited_by_user"]=True
            edited.append(ni)
    st.info(f"☑ {accepted} kabul • ☒ {rejected} reddet • ✏️ {manual} manuel düzenleme")
    return edited

st.markdown("""
<div class="opus-hero">
  <div class="opus-kicker">OPUS • PRIVATE COUNSEL SYSTEM</div>
  <h1 style="margin:.2rem 0 .35rem 0">OPUS LEGAL INTELLIGENCE</h1>
  <div style="color:#aaa59b">Contract intelligence • Negotiation strategy • Revision memory • Word redline</div>
</div>
""", unsafe_allow_html=True)

selected_module = st.radio(
    "Modül",
    ["Sözleşmeler", "Arabuluculuk", "KVKK", "Dava Dosyaları"],
    horizontal=True,
    label_visibility="collapsed"
)

cols = st.columns(4)
module_data = [
    ("Sözleşmeler", "Analyse • Negotiate • Revise", selected_module == "Sözleşmeler"),
    ("Arabuluculuk", "Prepare • Generate", selected_module == "Arabuluculuk"),
    ("KVKK", "Privacy • Compliance", selected_module == "KVKK"),
    ("Dava Dosyaları", "Litigation Workspace", selected_module == "Dava Dosyaları"),
]
for col, (title, sub, active) in zip(cols, module_data):
    with col:
        cls = "opus-module active" if active else "opus-module"
        st.markdown(
            f'<div class="{cls}"><div class="opus-module-title">{title}</div>'
            f'<div class="opus-module-sub">{sub if active else "Yakında • " + sub}</div></div>',
            unsafe_allow_html=True
        )

st.divider()
if selected_module == "Arabuluculuk":
    render_mediation()
    st.stop()
elif selected_module == "KVKK":
    st.header("KVKK")
    st.info("Bu modül sonraki aşamada aktif edilecek.")
    st.stop()
elif selected_module == "Dava Dosyaları":
    st.header("Dava Dosyaları")
    st.info("Bu modül sonraki aşamada aktif edilecek.")
    st.stop()

st.header("Yeni Sözleşme Analizi")

c1,c2,c3 = st.columns(3)
with c1:
    contract_type = st.selectbox("Sözleşme türü", ["Oyuncu Sözleşmesi"])
with c2:
    project_type = st.selectbox("Proje türü", ["Ana Akım TV", "Dijital", "Sinema"])
with c3:
    negotiation_power = st.select_slider(
        "Pazarlık gücü", ["Düşük","Orta","Yüksek","Çok Yüksek"], value="Orta"
    )

initial_note = st.text_area("İlk Not / Ajans Notu", placeholder="Bu dosyaya özgü bilgi veya talimatı yaz. Örn. ücret tamam; münhasırlık önemli.", height=90)
uploaded = st.file_uploader("Sözleşmeyi yükle", type=["docx","pdf"])
if project_type != "Ana Akım TV":
    st.info("Aktif profil şu an ACTOR_TV_MAINSTREAM.")

if uploaded:
    try:
        original_bytes = uploaded.getvalue()
        text = extract_text(uploaded)
        st.session_state["contract_text"] = text
        st.session_state["original_bytes"] = original_bytes
        st.session_state["uploaded_name"] = uploaded.name
        st.success(f"{uploaded.name} okundu • {len(text):,} karakter")

        with st.expander("Belgeden çıkarılan metni kontrol et"):
            st.text_area("Belge metni", text[:20000], height=220)

        if not get_api_key():
            st.warning("OPENAI_API_KEY tanımlı değil.")
        elif project_type == "Ana Akım TV":
            if st.button("⚖️ OLI Analizini Çalıştır", type="primary", use_container_width=True):
                with st.spinner("OLI 30 kontrol noktasını inceliyor..."):
                    result = analyse_contract(text, negotiation_power, initial_note)
                    st.session_state["oli_result"] = result
                    st.session_state.pop("revision_drafts", None)
    except Exception as e:
        st.error(f"Belge/analiz hatası: {e}")

result = st.session_state.get("oli_result")
if result:
    st.divider()
    st.header("Analiz Sonucu")

    findings = result.get("findings", [])
    a,b,c,d = st.columns(4)
    a.metric("Genel Risk", result.get("overall_risk","-"))
    b.metric("Kritik", sum(f.get("status")=="RED" for f in findings))
    c.metric("Müzakere", sum(f.get("status")=="ORANGE" for f in findings))
    d.metric("Korunacak", sum(f.get("status")=="GREEN" for f in findings))

    st.subheader("Yönetici Özeti")
    st.write(result.get("executive_summary",""))

    st.subheader("Masaya Getirilecek Konular")
    for item in result.get("top_negotiation_points", [])[:5]:
        st.write("• " + item)

    st.subheader("İlk İnceleme Notu")
    st.caption("Ajansa gönderilecek kısa müdahale özeti.")
    if st.button("📝 İlk İnceleme Notunu Hazırla", use_container_width=True):
        with st.spinner("Kısa bilgilendirme notu hazırlanıyor..."):
            try:
                st.session_state["initial_review_note"] = build_initial_review_note(result)
            except Exception as e:
                st.error(f"İlk inceleme notu hazırlanamadı: {e}")
    note_items = st.session_state.get("initial_review_note", {}).get("note_items", [])
    if note_items:
        note_text = "\n".join(
            f"{x.get('reference','')} {x.get('title','')}: {x.get('note','')}".strip()
            for x in note_items
        )
        st.text_area("Kopyalanabilir İlk İnceleme Notu", note_text, height=180)

    st.subheader("OLI Ek Bulgular")
    st.caption("30 Opus kuralı dışında kalan olası riskler. Bunlar otomatik revizyona alınmaz.")
    if st.button("🔎 30 Kural Dışı Riskleri Tara", use_container_width=True):
        with st.spinner("Sözleşme ikinci katman risk taramasından geçiyor..."):
            try:
                st.session_state["extra_risks"] = analyse_extra_risks(st.session_state["contract_text"])
            except Exception as e:
                st.error(f"Ek risk taraması çalıştırılamadı: {e}")
    extras = st.session_state.get("extra_risks", {}).get("extra_findings", [])
    if extras:
        for x in extras:
            with st.expander(f"🧭 {x.get('risk','-')} — {x.get('title','Ek bulgu')}"):
                st.write("**Referans:**", x.get("reference","-"))
                st.write("**Değerlendirme:**", x.get("assessment",""))
                st.write("**Önerilen aksiyon:**", x.get("suggested_action",""))
    elif "extra_risks" in st.session_state:
        st.success("30 kural dışında ayrıca anlamlı bir risk tespit edilmedi.")

    st.subheader("30 Kural Analizi")
    order={"RED":0,"ORANGE":1,"YELLOW":2,"GREEN":3,"NOT_FOUND":4}
    findings=sorted(findings,key=lambda f:order.get(f.get("status"),9))

    selectable=[]
    for f in findings:
        with st.expander(f"{icon(f.get('status'))} {f.get('rule_id','')} — {f.get('title','')}"):
            st.caption(
                f"Durum: {f.get('status','-')} • Müzakere: {f.get('negotiation_priority','-')} "
                f"• Güven: {f.get('confidence','-')}"
            )
            st.write("**Referans:**", f.get("contract_reference","-"))
            st.write("**Mevcut hüküm:**", f.get("clause_excerpt","-"))
            st.write("**OLI değerlendirmesi:**", f.get("assessment",""))
            lib = revision_context_for(f.get("rule_id"))
            if lib:
                st.write("**Opus Revision Library:**", lib.get("preferred_drafting",""))
            st.write("**İlk revizyon yaklaşımı:**", f.get("recommended_revision",""))

            default = f.get("status") in ("RED","ORANGE")
            if f.get("status") not in ("GREEN",):
                chosen = st.checkbox(
                    "Word revizyonuna dahil et",
                    value=default,
                    key=f"pick_{f.get('rule_id')}"
                )
                if chosen:
                    selectable.append(f)

    if uploaded and uploaded.name.lower().endswith(".docx"):
        st.divider()
        st.header("Word Revizyon Motoru")
        st.caption("OLI mevcut hükümleri revize eder; eksik koruyucu hükümleri uygun bölüme ekler. Word biçimi mümkün olduğunca korunur ve değişiklik yazarı Av. Onur Güneş olarak işlenir.")

        st.subheader("Word İşaretleri")
        st.caption("🟨 boş/doldurulacak alan • 🟧 risk/dikkat • 🟦 Rule Library'de anlamsal karşılığı olmayan yeni hüküm")
        if st.button("🎨 Sarı / Turuncu / Mavi İşaretleri Hazırla", use_container_width=True):
            with st.spinner("Word işaretleri belirleniyor..."):
                try:
                    st.session_state["word_flags"] = classify_word_flags(
                        st.session_state["contract_text"],
                        result,
                        st.session_state.get("extra_risks")
                    ).get("flags", [])
                except Exception as e:
                    st.error(f"Word işaretleri hazırlanamadı: {e}")
        flags = st.session_state.get("word_flags", [])
        if flags:
            blue = sum(x.get("color") == "blue" for x in flags)
            orange = sum(x.get("color") == "orange" for x in flags)
            st.info(f"🟧 {orange} risk/dikkat • 🟦 {blue} yeni/öğrenilmemiş hüküm bulundu.")

        if st.button("✍️ Revizyon Metinlerini Hazırla", use_container_width=True):
            if not selectable:
                st.warning("En az bir bulguyu Word revizyonuna dahil et.")
            else:
                with st.spinner("Opus revizyon kalıpları sözleşmeye uyarlanıyor..."):
                    drafts = build_revision_drafts(
                        st.session_state["contract_text"],
                        selectable,
                        negotiation_power
                    )
                    st.session_state["revision_drafts"] = drafts.get("revisions", [])

        drafts = st.session_state.get("revision_drafts", [])
        edited=[]
        if drafts:
            edited = render_revision_preview(
                drafts,
                st.session_state.get("contract_text", "")
            )
            st.session_state["approved_revision_plan"] = [dict(x) for x in edited]

            if st.button("📄 Onaylanan Revizyonlarla Word Oluştur", type="primary", use_container_width=True):
                plan = st.session_state.get("approved_revision_plan", [])
                if not plan:
                    st.warning("Word oluşturmak için en az bir revizyonu kabul et.")
                else:
                    try:
                        with st.spinner("Onaylanan revizyonlar Word'e uygulanıyor..."):
                            revised_bytes, applied, skipped, placeholder_count, flag_stats = apply_revisions_to_docx(
                                st.session_state["original_bytes"],
                                plan,
                                author="Av. Onur Güneş",
                                flags=st.session_state.get("word_flags", [])
                            )
                        st.session_state["revised_docx"] = revised_bytes
                        st.session_state["applied_revisions"] = applied
                        st.session_state["skipped_revisions"] = skipped
                        st.session_state["placeholder_count"] = placeholder_count
                        st.session_state["flag_stats"] = flag_stats
                        st.success(f"Word hazır. {len(applied)} revizyon uygulandı.")
                    except Exception as e:
                        st.error(f"Word oluşturulamadı: {e}")

            if st.session_state.get("revised_docx"):
                name = Path(st.session_state.get("uploaded_name","sozlesme.docx")).stem
                today_tr = datetime.now(ZoneInfo("Europe/Istanbul")).strftime("%d.%m.%Y")
                st.success(
                    f"{len(st.session_state.get('applied_revisions',[]))} revizyon Word'e işlendi."
                )
                skipped = st.session_state.get("skipped_revisions", [])
                if skipped:
                    st.warning(f"{len(skipped)} revizyon eşleşme bulunamadığı için uygulanamadı.")
                ph = st.session_state.get("placeholder_count",0)
                if ph:
                    st.info(f"🟨 {ph} doldurulması gereken alan sarı ile işaretlendi.")
                fs = st.session_state.get("flag_stats", {})
                if fs.get("orange") or fs.get("blue"):
                    st.info(f"🟧 {fs.get('orange',0)} risk/dikkat • 🟦 {fs.get('blue',0)} yeni/öğrenilmemiş hüküm Word'de işaretlendi.")
                st.download_button(
                    "⬇️ Revize Word'ü İndir",
                    data=st.session_state["revised_docx"],
                    file_name=f"{name} - {today_tr} REVİZE.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
    elif uploaded:
        st.info("Word üzerinde revizyon özelliği şu an yalnız .docx dosyalarında aktif.")


st.divider()
st.header("Karşı Taraf Dönüşü")
st.caption("Bizim gönderdiğimiz revize Word ile karşı taraftan dönen Word'ü karşılaştırır ve ajans bilgilendirme notunu çıkarır.")
returned_upload = st.file_uploader("Karşı taraftan dönen Word", type=["docx"], key="counterparty_return_upload")
if returned_upload and st.session_state.get("revised_docx"):
    if st.button("🔄 Karşı Taraf Revizyonunu Karşılaştır", use_container_width=True):
        try:
            our_text = read_docx(st.session_state["revised_docx"])
            returned_text = read_docx(returned_upload.getvalue())
            with st.spinner("Müzakere farkları sınıflandırılıyor..."):
                st.session_state["counterparty_review"] = compare_counterparty_return(our_text, returned_text)
        except Exception as e:
            st.error(f"Karşılaştırma yapılamadı: {e}")

    cp_items = st.session_state.get("counterparty_review", {}).get("items", [])
    if cp_items:
        icons = {"ACCEPTED":"🟢","PARTIAL":"🟡","REJECTED":"🔴","NEW":"🔵"}
        lines=[]
        for x in cp_items:
            label = icons.get(x.get("status"),"⚪")
            st.write(f"{label} **{x.get('reference','')} {x.get('title','')}** — {x.get('note','')}")
            lines.append(f"{x.get('reference','')} {x.get('note','')}".strip())
        st.text_area("Revizyon Dönüş Notu", "\n".join(lines), height=180)

        st.text_area(
            "Ajans Geri Bildirimi",
            key="agency_feedback",
            placeholder="Örn. 5.4 kabul. KDV'de ısrar. Cezai şart maksimum 3 olabilir.",
            height=100
        )
        st.caption("Ajans geri bildirimi dosyaya özgüdür; Rule Library veya Madde Bankası'na otomatik eklenmez.")




st.divider()
st.header("🧠 Toplu Öğrenme")
st.caption("Birden fazla ham ve revize Word yükle. OLI çiftleri eşleştirir, değişiklikleri kümeler ve senin onayladığın drafting davranışlarını topluca öğrenir.")

bc1,bc2=st.columns(2)
batch_raw=bc1.file_uploader("Ham sözleşmeler",type=["docx"],accept_multiple_files=True,key="batch_raw_v56")
batch_rev=bc2.file_uploader("Senin revize ettiğin sözleşmeler",type=["docx"],accept_multiple_files=True,key="batch_rev_v56")

if batch_raw and batch_rev and st.button("🔗 Sözleşmeleri Eşleştir",use_container_width=True):
    try:
        pairs,ur,uv=match_batch_documents(
            [(f.name,f.getvalue()) for f in batch_raw],
            [(f.name,f.getvalue()) for f in batch_rev]
        )
        st.session_state["batch_pairs_v56"]=pairs
        st.session_state["batch_unmatched_v56"]=(ur,uv)
    except Exception as e:
        st.error(f"Eşleştirme yapılamadı: {e}")

pairs=st.session_state.get("batch_pairs_v56",[])
if pairs:
    st.subheader("Eşleşmeler")
    confident=sum(p["status"]=="CONFIDENT" for p in pairs)
    review=len(pairs)-confident
    st.info(f"{len(pairs)} çift bulundu · {confident} güçlü eşleşme · {review} manuel kontrol")
    for i,p in enumerate(pairs):
        icon="🟢" if p["status"]=="CONFIDENT" else "🟡"
        st.write(f"{icon} **{p['raw_name']}** ↔ **{p['revised_name']}** · eşleşme {p['match_score']}")
    ur,uv=st.session_state.get("batch_unmatched_v56",([],[]))
    if ur or uv:
        st.warning("Eşleşmeyen dosyalar: "+", ".join(ur+uv))
    if st.button("🔬 Tüm Çiftlerden Öğrenme Özeti Çıkar",use_container_width=True):
        try:
            cs=batch_candidates(pairs)
            clusters=cluster_candidates(cs)
            st.session_state["batch_candidates_v56"]=cs
            st.session_state["batch_clusters_v56"]=clusters
        except Exception as e:
            st.error(f"Toplu öğrenme analizi yapılamadı: {e}")

clusters=st.session_state.get("batch_clusters_v56",[])
allc=st.session_state.get("batch_candidates_v56",[])
if clusters:
    st.subheader("Toplu Öğrenme Özeti")
    st.caption(f"{len(allc)} tekil değişiklik, {len(clusters)} davranış kümesine ayrıldı. Küme onayı, içindeki örnekleri topluca öğrenir.")
    cluster_decisions={}
    for i,c in enumerate(clusters):
        with st.expander(f"{c['topic']} — {c['edit_style']} · {c['count']} örnek · güven {c['confidence']}",expanded=(i<4)):
            for ex in c["examples"][:3]:
                st.markdown(f"**{ex.get('pair_raw_name','')} → {ex.get('pair_revised_name','')}**")
                st.caption("Önce: "+(ex.get("original_text") or "—")[:350])
                st.caption("Nihai: "+(ex.get("final_text") or "—")[:350])
            cc1,cc2,cc3=st.columns([1,1.4,1])
            ok=cc1.checkbox("✓ Kümeyi Öğren",key=f"bc_ok_{i}")
            topic=cc2.text_input("Konu",value=c["topic"],key=f"bc_topic_{i}")
            scope=cc3.selectbox("Kapsam",["GENERAL","FILE_ONLY"],key=f"bc_scope_{i}")
            cluster_decisions[c["cluster_id"]]={"approve":ok,"topic":topic,"scope":scope}
    if st.button("💾 Onaylanan Kümeleri Hafızaya Kaydet",type="primary",use_container_width=True):
        decisions={}
        for c in clusters:
            d=cluster_decisions.get(c["cluster_id"],{})
            for cid in c["candidate_ids"]:
                decisions[cid]={"approve":d.get("approve",False),"topic":d.get("topic",c["topic"]),"scope":d.get("scope","GENERAL")}
        try:
            added,mem=approve_candidates(allc,decisions)
            st.success(f"{added} tekil emsal öğrenildi. Toplam onaylı kayıt: {sum(1 for r in mem['records'] if r.get('approved'))}")
        except Exception as e:
            st.error(f"Toplu öğrenme kaydedilemedi: {e}")

st.divider()
st.header("Nihai Revizyondan Öğren — Learning Engine v1")
st.caption("Geçmiş dosyalarda ilk gelen Word + senin nihai revize Word'ünü karşılaştırır. Hiçbir kayıt sen onaylamadan kalıcı öğrenme hafızasına girmez.")

lc1,lc2=st.columns(2)
learn_original=lc1.file_uploader("İlk gelen sözleşme (.docx)",type=["docx"],key="learn_original_v1")
learn_final=lc2.file_uploader("Nihai / senin revize ettiğin sözleşme (.docx)",type=["docx"],key="learn_final_v1")

if learn_original and learn_final and st.button("🧠 Öğrenme Adaylarını Çıkar",use_container_width=True):
    try:
        candidates=build_candidates(
            learn_original.getvalue(),
            learn_final.getvalue(),
            source_name=learn_final.name
        )
        st.session_state["learning_candidates_v1"]=candidates
    except Exception as e:
        st.error(f"Öğrenme adayları çıkarılamadı: {e}")

cands=st.session_state.get("learning_candidates_v1",[])
if cands:
    st.info(f"{len(cands)} değişiklik/adisyon öğrenme adayı bulundu. Yalnız onayladıkların kaydedilecek.")
    decisions={}
    for i,c in enumerate(cands):
        with st.expander(f"{i+1}. {c.get('edit_style')} — {c.get('final_text','')[:90]}",expanded=(i<3)):
            st.markdown("**ÖNCE**")
            st.text_area("Önce",c.get("original_text") or "—",height=100,disabled=True,key=f"l_old_{i}",label_visibility="collapsed")
            st.markdown("**NİHAİ TERCİH**")
            st.text_area("Nihai",c.get("final_text") or "—",height=110,disabled=True,key=f"l_new_{i}",label_visibility="collapsed")
            cc1,cc2,cc3=st.columns([1,1.4,1.2])
            approve=cc1.checkbox("✓ Öğren",key=f"l_ok_{i}")
            topic=cc2.text_input("Konu",value=c.get("topic","SINIFLANDIRILMADI"),key=f"l_topic_{i}")
            scope=cc3.selectbox("Kapsam",["GENERAL","FILE_ONLY"],key=f"l_scope_{i}")
            decisions[c["candidate_id"]]={"approve":approve,"topic":topic,"scope":scope}
    if st.button("💾 Onaylananları Öğrenme Hafızasına Kaydet",type="primary",use_container_width=True):
        try:
            added,mem=approve_candidates(cands,decisions)
            st.success(f"{added} öğrenme kaydı kaydedildi. Toplam onaylı kayıt: {sum(1 for r in mem['records'] if r.get('approved'))}")
            st.session_state.pop("learning_candidates_v1",None)
        except Exception as e:
            st.error(f"Öğrenme hafızası kaydedilemedi: {e}")

mem=load_memory()
approved=[r for r in mem.get("records",[]) if r.get("approved")]
with st.expander(f"📚 Öğrenme Hafızası ({len(approved)} onaylı kayıt)"):
    if not approved:
        st.caption("Henüz onaylı öğrenme kaydı yok.")
    else:
        for r in approved[-20:][::-1]:
            st.write(f"**{r.get('topic')}** · {r.get('edit_style')} · güven {r.get('confidence')} · {r.get('precedent_count',1)} örnek")
            st.caption((r.get("final_text") or "")[:300])

st.divider()
st.header("Nihai Revizyondan Öğren")
st.caption("OLI çıktısını sen son kez revize ettikten sonra buraya yükle. Sistem farkları öğrenme adayı olarak çıkarır; hiçbir şeyi otomatik olarak Revision Library'ye eklemez.")
final_upload = st.file_uploader("Av. Onur Güneş nihai revize Word", type=["docx"], key="final_revision_upload")
if final_upload and st.session_state.get("revised_docx"):
    if st.button("🧠 OLI Revizyonu ile Karşılaştır", use_container_width=True):
        try:
            oli_text = read_docx(st.session_state["revised_docx"])
            final_text = read_docx(final_upload.getvalue())
            with st.spinner("Nihai revizyon tercihleri karşılaştırılıyor..."):
                st.session_state["learning_review"] = compare_final_revision(oli_text, final_text)
        except Exception as e:
            st.error(f"Karşılaştırma yapılamadı: {e}")

    candidates = st.session_state.get("learning_review", {}).get("learning_candidates", [])
    for c in candidates:
        with st.expander(f"🧠 {c.get('title','Revizyon tercihi')}"):
            st.write("**OLI:**", c.get("oli_position",""))
            st.write("**Nihai tercih:**", c.get("final_position",""))
            st.write("**Fark:**", c.get("difference",""))
            st.caption(f"Öneri: {c.get('recommendation')} • Güven: {c.get('confidence')}")
            st.radio(
                "Bu öğrenme ne olsun?",
                ["Henüz öğrenme", "Revision Library adayı", "Bu dosyaya özgü"],
                index=0,
                key=f"learn_{c.get('title','')}_{c.get('difference','')[:20]}"
            )


st.divider()
with st.expander("OLI Bilgi Tabanı"):
    st.write(f"**Rule Library:** {len(RULES)} kontrol noktası")
    st.write(f"**Revision Library:** {len(REVISION_LIBRARY.get('entries',[]))} doğrulanmış revizyon kalıbı")
    st.write(f"**Madde Bankası:** {len(CLAUSE_BANK.get('entries',[]))} hazır Opus cümlesi")

st.caption("OLI • Sözleşmeler v0.5.6 Batch Learning + Arabuluculuk v1.3.1 • Prototip. Gerçek müvekkil belgeleri için erişim, veri güvenliği, saklama ve meslek sırrı mimarisi ayrıca tamamlanmalıdır.")
