from __future__ import annotations

import json
import re
import sqlite3
from datetime import date
from io import BytesIO
from pathlib import Path

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from pypdf import PdfReader
from openai import OpenAI

from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak


MEDIATORS = {
    "Av. Arb. Onur GÜNEŞ": {
        "name": "Onur GÜNEŞ",
        "full_name": "Av. Arb. Onur GÜNEŞ",
        "registry": "11706",
        "address": "Asmalı Mescit Mah. Tünel Meydanı Tünel Geçidi İş Merkezi B Blok D:416 Beyoğlu/İSTANBUL",
    },
    "Av. Arb. Serap TAŞDELEN GÜNEŞ": {
        "name": "Serap TAŞDELEN GÜNEŞ",
        "full_name": "Av. Arb. Serap TAŞDELEN GÜNEŞ",
        "registry": "37393",
        "address": "Osmanağa Mah. Serasker Cad. Pavlonya Sok. No:25/8 Kadıköy/İSTANBUL",
    },
}


INFO_COMMON = """Taraflara arabuluculuk sürecinin iradi olduğu, arabuluculuk sürecinde her iki tarafın da eşit haklara sahip olduğu, taraflarca aksi kararlaştırılmadıkça arabulucunun arabuluculuk faaliyeti çerçevesinde kendisine sunulan veya diğer bir şekilde elde ettiği bilgi ve belgeler ile diğer kayıtları gizli tutmakla yükümlü olduğu ve tarafların ve görüşmelere katılan diğer kişilerin de bu konudaki gizliliğe uymak zorunda olduğu, tarafların, arabulucunun veya arabuluculuğa katılanlar dahil üçüncü bir kişinin, uyuşmazlıkla ilgili hukuk davası açıldığında veya tahkim yoluna başvurulduğunda; tarafların arabuluculuk sürecine katılma isteğini, arabuluculuk sürecinde taraflarca ileri sürülen görüşlerini, önerileri veya herhangi bir olay ve iddianın kabulünü ve sadece arabuluculuk faaliyeti dolayısıyla hazırlanan belgeleri delil olarak ileri süremeyeceği ve bunlar hakkında tanıklık yapamayacağına dair arabuluculuğun temel ilkeleri hakkında bilgi verilmiştir.

Taraflara arabulucunun görevini özenle, tarafsız bir biçimde ve şahsen yerine getireceği, arabulucunun taraflar arasında eşitliği gözetmekle yükümlü olduğu, arabuluculuk müzakerelerine tarafların bizzat, kanuni temsilcileri veya vekaletnamesinde özel yetki bulunan avukatları aracılığıyla katılabileceği, arabuluculuk sürecinde arabulucunun rolünün, hakim veya hakem olmadığı, kimin haklı ya da haksız olduğu konusunda karar vermeyeceği, yargısal bir yetkinin kullanımı olarak sadece hakim tarafından yapılabilecek işlemleri yapamayacağı, taraflara hukuki tavsiyelerde bulunamayacağı, tarafların çözüm üretemediklerinin ortaya çıkması halinde arabulucunun bir çözüm önerisinde bulunabileceği, yaşanılan uyuşmazlık ile ilgili çözüm seçeneklerini üreterek bir anlaşmaya ulaşabilmelerinde taraflara yardımcı olacak iletişimin ortamını sağlayacağı, bilgileri dahilinde taraflarla ayrı ayrı veya birlikte görüşebileceği ve iletişim kurabileceği, arabulucu olarak tarafsız bir konumda olduğu, arabuluculuk sürecinin sonunda her iki tarafın da kabul edeceği bir anlaşmaya varılamaması hâlinde açılabilecek olası bir davada, daha sonra avukat olarak görev üstlenemeyeceği, arabuluculuk bürosuna başvurulmasından son tutanağın düzenlendiği tarihe kadar geçen sürede zamanaşımının duracağı ve hak düşürücü sürenin işlemeyeceği, arabuluculuk sürecinin sonunda her iki tarafın da kabul edeceği bir anlaşmaya varılamaması hâlinde yargı organlarına başvuru haklarının bulunduğu hususları hakkında bilgi verilmiştir.

Taraflara arabuluculuk sürecinde düzenlenecek oturum tutanaklarına ve sürecin sonunda düzenlenecek son tutanağa, oturumların ve faaliyetin sonuçlanması ile arabulucunun son çözüm önerisi dışında hangi hususların yazılacağına tarafların karar vereceği; bununla beraber, bu hususta tarafların birlikte karar verememesi halinde son tutanağın içeriğinin tarafların karşılıklı teklif ve kabulleri dahil olmak üzere arabulucu tarafından düzenleneceği, arabuluculuk sürecinin sonunda varılan anlaşmanın kapsamının taraflarca belirleneceği, anlaşma belgesi düzenlenmesi hâlinde bu belgenin taraflar veya avukatları ve arabulucu tarafından imzalanacağı, tarafların bu anlaşma belgesinin icra edilebilirliğine ilişkin mahkemeden şerh verilmesini talep edebileceği ve bu şerhi içeren anlaşmanın ilâm niteliğinde belge sayılacağı, taraflar ve avukatları ile arabulucunun birlikte imzaladıkları anlaşma belgesinin icra edilebilirlik şerhi aranmaksızın ilâm niteliğinde belge sayılacağı, arabuluculuk faaliyeti sonunda anlaşmaya varılması hâlinde üzerinde anlaşılan hususlar hakkında taraflarca dava açılamayacağı hususları hakkında bilgi verilmiştir.

Taraflara arabuluculuk faaliyeti sonunda anlaşmaları halinde, arabuluculuk ücretinin, aksi kararlaştırılmadıkça taraflarca eşit şekilde karşılanacağı, ücretin tarifenin birinci kısmında belirlenen iki saatlik ücret tutarından az olamayacağı, arabuluculuk faaliyeti sonunda iki saatten az süren görüşmeler sonunda tarafların anlaşamamaları hâllerinde iki saatlik ücret tutarının tarifenin birinci kısmına göre ve aksi kararlaştırılmadıkça taraflarca eşit şekilde karşılanacağı konusunda bilgi verildi.

Taraflara 6698 sayılı Kanunun (KVKK) 10. maddesi gereğince kişisel verilerin; hangi amaçla işleneceği, kimlere ve hangi amaçla aktarılabileceği, veri toplamanın yöntemi ve hukuki sebebi hakkında bilgi verildi. Devamla yasanın 11. maddesi gereğince veri sorumlusuna (arabulucu) başvurarak kendileriyle ilgili kişisel verilerin; işlenip işlenmediğini öğrenme, işlenmişse buna ilişkin bilgi talep etme, işlenme amacını ve bunların amacına uygun kullanılıp kullanılmadığını öğrenme, yurt içinde veya yurt dışında kişisel verilerin aktarıldığı üçüncü kişileri bilme, verilerin eksik veya yanlış işlenmiş olması hâlinde bunların düzeltilmesini isteme, yasal şartlar çerçevesinde kişisel verilerin silinmesini veya yok edilmesini isteme, işlenen verilerin münhasıran otomatik sistemler vasıtasıyla analiz edilmesi suretiyle kişinin kendisi aleyhine bir sonucun ortaya çıkmasına itiraz etme ve kanuna aykırı olarak işlenmesi sebebiyle zarara uğraması hâlinde zararın giderilmesini talep etme haklarına sahip oldukları bildirilmiştir.

Taraflara, arabuluculuk sürecinde olabildiğince açık ve dürüst olunmasının ve işbirliği hâlinde hareket edilmesinin önemi vurgulanmış ve arabuluculuk sürecinde belirtilen kurallara uymayı kabul edip etmedikleri kendilerine sorulmuştur.

Taraflar söz alarak; arabuluculuğun temel ilkelerini, arabuluculuk sürecini ve arabuluculuk süreci sonunda hazırlanan arabuluculuk son tutanağının ve anlaşma belgesinin hukuki ve mali yönlerden bütün sonuçlarını anladıklarını ve kurallara uymayı kabul ettiklerini beyan etmişlerdir."""


def db_path():
    # Community Cloud storage is prototype-only; this persists while the instance filesystem survives.
    return Path(__file__).with_name("mediation_parties.db")


def init_db():
    conn = sqlite3.connect(db_path())
    conn.execute("""CREATE TABLE IF NOT EXISTS parties (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        party_type TEXT,
        name TEXT NOT NULL,
        identity_no TEXT,
        tax_no TEXT,
        address TEXT,
        attorney TEXT,
        attorney_bar TEXT,
        UNIQUE(name, identity_no, tax_no)
    )""")
    conn.commit()
    return conn


def search_parties(prefix: str):
    prefix = (prefix or "").strip()
    if len(prefix) < 2:
        return []
    conn = init_db()
    rows = conn.execute(
        """SELECT id, party_type, name, identity_no, tax_no, address, attorney, attorney_bar
           FROM parties WHERE lower(name) LIKE lower(?) ORDER BY name LIMIT 20""",
        (prefix + "%",)
    ).fetchall()
    conn.close()
    keys = ["id","party_type","name","identity_no","tax_no","address","attorney","attorney_bar"]
    return [dict(zip(keys, r)) for r in rows]


def save_party(p):
    if not (p.get("name") or "").strip():
        return
    conn = init_db()
    conn.execute(
        """INSERT OR REPLACE INTO parties
        (id, party_type, name, identity_no, tax_no, address, attorney, attorney_bar)
        VALUES (
          COALESCE((SELECT id FROM parties WHERE name=? AND COALESCE(identity_no,'')=? AND COALESCE(tax_no,'')=?), NULL),
          ?,?,?,?,?,?,?
        )""",
        (
            p.get("name",""), p.get("identity_no",""), p.get("tax_no",""),
            p.get("party_type","Gerçek Kişi"), p.get("name",""), p.get("identity_no",""),
            p.get("tax_no",""), p.get("address",""), p.get("attorney",""), p.get("attorney_bar","")
        )
    )
    conn.commit()
    conn.close()


def fmt_date(v):
    if not v:
        return "[TARİH]"
    if hasattr(v, "strftime"):
        return v.strftime("%d.%m.%Y")
    return str(v)


def party_block(i, p):
    label = "Adı Soyadı" if p.get("party_type") == "Gerçek Kişi" else "Unvanı"
    lines = [f"TARAF {i} :", f"{label} : {p.get('name') or '[AD / UNVAN]'}"]
    if p.get("identity_no"):
        lines.append(f"T.C. Kimlik No : {p['identity_no']}")
    if p.get("tax_no"):
        lines.append(f"Vergi Numarası : {p['tax_no']}")
    if p.get("attorney"):
        bar = f" ({p['attorney_bar']})" if p.get("attorney_bar") else ""
        lines.append(f"Vekili : {p['attorney']}{bar}")
    if p.get("address"):
        lines.append(f"Adres : {p['address']}")
    return "\n".join(lines)


def all_parties_header(parties):
    return "\n".join(party_block(i+1, p) for i,p in enumerate(parties) if (p.get("name") or "").strip())


def signatures(parties, mediator):
    lines = ["İMZALAR"]
    for p in parties:
        if not (p.get("name") or "").strip():
            continue
        ident = f" (TC: {p['identity_no']})" if p.get("identity_no") else ""
        lines.append(f"{p['name']}{ident}")
        if p.get("attorney"):
            bar = f" ({p['attorney_bar']})" if p.get("attorney_bar") else ""
            lines.append(f"Vekili {p['attorney']}{bar}")
    lines.append(f"Arabulucu {mediator['name']} ({mediator['registry']})")
    return "\n".join(lines)


def dispute_label(kind):
    return {
        "İşçilik Alacağı": "İşçi - İşveren Uyuşmazlığı (İşçilik Alacağı)",
        "Kira Tespit": "Kira Parasının Tespiti",
        "Kiralananın Tahliyesi": "Kiralananın Tahliyesi",
        "Kira Tespit + Kiralananın Tahliyesi": "Kira Parasının Tespiti ve Kiralananın Tahliyesi",
    }[kind]


def info_intro(kind, mediator):
    if kind == "İşçilik Alacağı":
        desc = "işçi işveren uyuşmazlığından kaynaklı işçilik alacağı konulu uyuşmazlığın"
    else:
        desc = "kira ilişkisinden kaynaklı uyuşmazlığın"
    return (
        f"Taraflar, {desc} çözümü için arabulucu olarak Arabuluculuk Daire Başkanlığı resmi siciline "
        f"kayıtlı {mediator['registry']} sicil numaralı {mediator['full_name']}'in atanmasına karar vermişlerdir."
    )


def mediator_header(m):
    return (
        "Arabulucunun :\n"
        f"Adı ve Soyadı : {m['name']}\n"
        f"Arabulucu Sicil Numarası : {m['registry']}\n"
        f"Adres : {m['address']}"
    )


def build_information_text(data):
    m, parties = data["mediator"], data["parties"]
    opening = fmt_date(data.get("opening_date"))
    return f"""İHTİYARİ ARABULUCULUKTA BİLGİLENDİRME VE BELİRLEME TUTANAĞI

{mediator_header(m)}

{all_parties_header(parties)}

{info_intro(data["dispute_kind"], m)}

{INFO_COMMON}

İşbu arabuluculuk bilgilendirme ve belirleme tutanağı taraflarca imza altına alınmıştır.
{opening}

{signatures(parties, m)}"""


def employment_agreement_terms(data):
    d = data["details"]
    p1 = data["parties"][0]["name"] if data["parties"] else "TARAF 1"
    p2 = data["parties"][1]["name"] if len(data["parties"]) > 1 else "TARAF 2"
    amount = d.get("amount") or "[TUTAR]"
    netgross = d.get("netgross") or "[NET/BRÜT]"
    receivable = d.get("receivable") or "[ALACAK KALEMİ]"
    payment_date = fmt_date(d.get("payment_date"))
    termination_date = fmt_date(d.get("termination_date"))
    method = d.get("payment_method") or "[ÖDEME YÖNTEMİ]"
    return f"""1. TARAF 1 {p1} ile TARAF 2 {p2}, iş sözleşmesinin {termination_date} tarihinde sona erdiği hususunda anlaşmışlardır.

2. TARAF 1'in {receivable} alacağı için {netgross} {amount} TL'nin TARAF 2 tarafından TARAF 1'e {payment_date} tarihinde {method} ödenmesi hususunda anlaşmışlardır.

3. Taraflar, yukarıda belirtilen ödemeler dışında iş sözleşmesinin feshi sebebiyle, işe iade, boşta geçen süre ücreti, işe başlatmama tazminatı, kötü niyet tazminatı, bireysel ve toplu iş sözleşmesinden kaynaklanan ücret ve tüm alacak ferileri ile hiçbir nam ve isim adı altında alacak ve geçmişe dönük maddi ve manevi hakları ile ilgili hiçbir hak ve taleplerinin bulunmadığı hususunda anlaşmışlardır.

4. Tarafların, yukarıda belirtilen yükümlülüklerini yerine getirmeleri halinde birbirlerinden herhangi bir hak ve alacağı kalmamış olacak ve taraflar birbirlerini gayri kabili rücu olarak ibra edilmiş sayacak, taraflar arasında çalışılan döneme ilişkin uyuşmazlık konusu ve dava konusu yapılabilecek herhangi bir ihtilaf da kalmamış olacaktır."""


def rent_agreement_terms(data):
    d = data["details"]
    p1 = data["parties"][0]["name"] if data["parties"] else "TARAF 1"
    p2 = data["parties"][1]["name"] if len(data["parties"]) > 1 else "TARAF 2"
    addr = d.get("property_address") or "[TAŞINMAZ ADRESİ]"
    lease_start = fmt_date(d.get("lease_start"))
    end_date = fmt_date(d.get("lease_end"))
    eviction = fmt_date(d.get("eviction_date"))
    lines = [
        f"1. Taraflar, TARAF 1 {p1} ile TARAF 2 {p2} arasında, {addr} adresinde bulunan taşınmaza ilişkin {lease_start} başlangıç tarihli kira sözleşmesinin, {end_date} tarihi itibarıyla sona ereceği ve bu tarih itibarıyla taraflar arasında ayrıca herhangi bir ihtara ve ihbara gerek olmaksızın karşılıklı anlaşma ile feshedilmiş sayılacağı hususunda anlaşmışlardır.",
        f"2. Taraflar, TARAF 2'nin Kiracı sıfatıyla kullandığı {addr} adresindeki taşınmazı, {eviction} tarihinde, kira sözleşmesi kapsamında teslim aldığı haliyle, boş, sağlam ve hasarsız olarak tahliye edip TARAF 1'e teslim edeceği hususunda anlaşmışlardır."
    ]
    idx = 3
    for period in d.get("rent_periods", []):
        if not any(period.values()):
            continue
        lines.append(
            f"{idx}. TARAF 2, {period.get('period') or '[DÖNEM]'} dönemine ait kira bedellerini "
            f"{period.get('payment_rule') or '[ÖDEME ŞEKLİ]'} olmak üzere, aylık net "
            f"{period.get('amount') or '[TUTAR]'} TL olarak TARAF 1'e ödeyecektir."
        )
        idx += 1
    lines.append(
        f"{idx}. TARAF 2'nin yukarıda belirtilen kira bedellerini süresinde ve eksiksiz olarak ödemesi halinde, "
        "TARAF 1'in belirtilen kira dönemlerine ilişkin ayrıca kira alacağı kalmayacağı hususunda anlaşmışlardır."
    )
    return "\n\n".join(lines)


def fee_clause(data, no):
    d = data["details"]
    fee = d.get("mediation_fee") or "[ARABULUCULUK ÜCRETİ]"
    payer = d.get("fee_payer") or "[ÖDEYECEK TARAF]"
    paydate = fmt_date(d.get("fee_payment_date"))
    account = d.get("fee_account") or "[ÖDEME HESABI / IBAN]"
    vat = " + KDV" if d.get("fee_plus_vat") else ""
    return (
        f"{no}. Arabuluculuk ücretinin tamamı olan {fee} TL{vat}, {payer} tarafından en geç "
        f"{paydate} tarihinde Arabulucu {data['mediator']['name']}'in {account} hesabına ödenmesi hususunda anlaşmışlardır."
    )


def build_agreement_text(data):
    m, parties = data["mediator"], data["parties"]
    start = fmt_date(data.get("opening_date"))
    agreement = fmt_date(data.get("agreement_date"))
    final = fmt_date(data.get("final_date"))
    kind = data["dispute_kind"]

    if kind == "İşçilik Alacağı":
        terms = employment_agreement_terms(data)
        fee_no = 5
        law = "6325 sayılı Hukuk Uyuşmazlıklarında Arabuluculuk Kanunu m. 17 ve 7036 sayılı İş Mahkemeleri Kanunu m. 3"
    else:
        terms = rent_agreement_terms(data)
        fee_no = len(re.findall(r"(?m)^\d+\.", terms)) + 1
        law = "6325 sayılı Hukuk Uyuşmazlıklarında Arabuluculuk Kanunu m. 17 ve m.18"

    return f"""HUKUK UYUŞMAZLIKLARINDA İHTİYARİ ARABULUCULUK ANLAŞMA BELGESİ
DOSYA NO : {data.get("file_no") or "[DOSYA NO]"}

{mediator_header(m)}

{all_parties_header(parties)}

Arabuluculuk Konusu Uyuşmazlık : {dispute_label(kind)}
Arabuluculuk Sürecinin Başladığı Tarih : {start}
Arabuluculuk Sürecinin Bittiği Tarih : {agreement}
Son Tutanağın Düzenlendiği Tarih : {final}
Arabuluculuk Sonucu : ANLAŞMA

Tarafların Arabulucu {m['name']}'in yardımıyla 6325 sayılı Hukuk Uyuşmazlıklarında Arabuluculuk Kanunu'nun 18. maddesi uyarınca anlaşmaya varmaları üzerine işbu anlaşma belgesi hazırlanmıştır.

Taraflara arabuluculuk anlaşma belgesinin hukuki ve mali yönlerden bütün sonuçları hakkında bilgi verildi. Taraflara arabuluculuğun temel ilkeleri, arabuluculuk süreci ve arabuluculuk süreci sonunda hazırlanan arabuluculuk son tutanağının hukuki ve mali yönlerden bütün sonuçları hakkında bilgi verildi.

Taraflar arabulucu huzurunda aralarındaki uyuşmazlığın çözümü konusunda özgür iradeleriyle aşağıdaki şartlarda anlaşmışlardır.

{terms}

{fee_clause(data, fee_no)}

İşbu anlaşma belgesi taraflarca {law} uyarınca hep birlikte imzalanmış ve taraflara birer örneği iletilmiştir. {agreement}

{signatures(parties, m)}"""


def build_final_text(data):
    m, parties = data["mediator"], data["parties"]
    kind = data["dispute_kind"]
    start = fmt_date(data.get("opening_date"))
    end = fmt_date(data.get("agreement_date"))
    final = fmt_date(data.get("final_date"))

    if kind == "İşçilik Alacağı":
        d = data["details"]
        receivable = d.get("receivable") or "[ALACAK KALEMİ]"
        summary = (
            f"Görüşmeler sırasında taraflar iş sözleşmesinin sonlandırıldığı, TARAF 1'in {receivable} dışındaki "
            "işçilik alacakları bakımından anlaşma şartlarının ayrıca düzenlendiği; bu kaleme ilişkin anlaşma belgesi "
            "doğrultusunda ödeme yapılacağı konusunda mutabık kalmışlardır.\n\n"
            "Taraflar arasında, işçi ile işveren ilişkisinden kaynaklı işçilik alacakları yönünden anlaşma sağlanmıştır."
        )
        law = "6325 sayılı Hukuk Uyuşmazlıklarında Arabuluculuk Kanunu m. 17 ve 7036 sayılı İş Mahkemeleri Kanunu m. 3"
    else:
        d = data["details"]
        addr = d.get("property_address") or "[TAŞINMAZ ADRESİ]"
        summary = (
            f"Taraflar, TARAF 1 Kiraya Veren ile TARAF 2 Kiracı arasında, {addr} adresinde bulunan taşınmaza ilişkin "
            "aylık kira bedeli ve ödeme koşulları, taraflar arasındaki kira sözleşmesinin sona erdirilmesi ve/veya "
            "kiralananın tahliyesi ile teslim tarihi hususlarında karşılıklı olarak mutabakata varmışlardır."
        )
        law = "6325 sayılı Hukuk Uyuşmazlıklarında Arabuluculuk Kanunu m. 17 ve m.18"

    return f"""HUKUK UYUŞMAZLIKLARINDA İHTİYARİ ARABULUCULUK SON TUTANAĞI
DOSYA NO : {data.get("file_no") or "[DOSYA NO]"}

{mediator_header(m)}

{all_parties_header(parties)}

Arabuluculuk Konusu Uyuşmazlık : {dispute_label(kind)}
Arabuluculuk Sürecinin Başladığı Tarih : {start}
Arabuluculuk Sürecinin Bittiği Tarih : {end}
Son Tutanağın Düzenlendiği Tarih : {final}
Arabuluculuk Sonucu : ANLAŞMA

Adı geçen taraflar ilk oturum gün ve saatinde toplantıya iştirak etmişlerdir.

Taraflara arabuluculuk anlaşma belgesinin hukuki ve mali yönlerden bütün sonuçları hakkında bilgi verildi. Taraflara arabuluculuğun temel ilkeleri, arabuluculuk süreci ve arabuluculuk süreci sonunda hazırlanan arabuluculuk son tutanağının hukuki ve mali yönlerden bütün sonuçları hakkında bilgi verildi.

{summary}

Taraflar arabuluculuk görüşmeleri sonucunda anlaşmışlardır. Taraflar arasında aynı tarihli anlaşma belgesi düzenlenmiş ve işbu son tutanak düzenlenerek görüşmeler son bulmuştur.

İşbu son tutanak {law} uyarınca hep birlikte imzalanmış ve taraflara birer örneği iletilmiştir. {final}

{signatures(parties, m)}"""


def text_to_docx(title_text: str):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = 18 * 36000
    sec.bottom_margin = 18 * 36000
    sec.left_margin = 20 * 36000
    sec.right_margin = 20 * 36000

    lines = title_text.splitlines()
    first = True
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        if first:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.bold = True
            run.font.size = Pt(12)
            first = False
        elif line.strip() in ("İMZALAR",):
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _pdf_font():
    try:
        import reportlab
        font = Path(reportlab.__file__).resolve().parent / "fonts" / "Vera.ttf"
        if font.exists():
            pdfmetrics.registerFont(TTFont("Vera", str(font)))
            return "Vera"
    except Exception:
        pass
    return "Helvetica"


def text_to_pdf(text: str):
    bio = BytesIO()
    font_name = _pdf_font()
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "OLINormal", parent=styles["Normal"], fontName=font_name,
        fontSize=9.5, leading=12.5, spaceAfter=5
    )
    title = ParagraphStyle(
        "OLITitle", parent=normal, fontSize=11, leading=14,
        alignment=TA_CENTER, spaceAfter=10
    )
    doc = SimpleDocTemplate(
        bio, pagesize=A4,
        rightMargin=18*mm, leftMargin=18*mm, topMargin=16*mm, bottomMargin=16*mm
    )
    story = []
    for i, block in enumerate(text.split("\n\n")):
        block = block.strip()
        if not block:
            continue
        safe = block.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;").replace("\n","<br/>")
        story.append(Paragraph(safe, title if i == 0 else normal))
        story.append(Spacer(1, 2*mm))
    doc.build(story)
    return bio.getvalue()


def extract_upload_text(uploaded):
    name = uploaded.name.lower()
    b = uploaded.getvalue()
    if name.endswith(".pdf"):
        reader = PdfReader(BytesIO(b))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    if name.endswith(".docx"):
        doc = Document(BytesIO(b))
        return "\n".join(p.text for p in doc.paragraphs)
    return ""


def ai_extract_exit_notice(text):
    try:
        key = st.secrets["OPENAI_API_KEY"]
        model = st.secrets.get("OPENAI_MODEL", "gpt-5.6-terra")
    except Exception:
        return None
    client = OpenAI(api_key=key)
    system = """İşten ayrılış/çıkış bildirgesindeki taraf bilgilerini çıkar.
Yalnız JSON:
{
 "employee":{"name":"","identity_no":"","address":""},
 "employer":{"name":"","tax_no":"","address":""},
 "termination_date":"GG.AA.YYYY veya boş"
}
Metinde olmayan bilgiyi uydurma."""
    resp = client.responses.create(
        model=model,
        input=[{"role":"system","content":system},{"role":"user","content":text[:30000]}]
    )
    raw = resp.output_text.strip()
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw)
    return json.loads(raw[raw.find("{"):raw.rfind("}")+1])


def _party_editor(i):
    st.markdown(f"#### Taraf {i}")
    prefix = st.text_input("Kayıtlı taraf ara", key=f"med_search_{i}", placeholder="İlk 2-3 harfi yaz...")
    matches = search_parties(prefix)
    if matches:
        labels = ["Seçiniz"] + [m["name"] for m in matches]
        choice = st.selectbox("Kayıtlı taraf", labels, key=f"med_match_{i}")
        if choice != "Seçiniz":
            m = next(x for x in matches if x["name"] == choice)
            if st.button("Kayıtlı tarafı getir", key=f"med_get_{i}"):
                st.session_state[f"med_p{i}_type"] = m["party_type"]
                st.session_state[f"med_p{i}_name"] = m["name"]
                st.session_state[f"med_p{i}_id"] = m["identity_no"] or ""
                st.session_state[f"med_p{i}_tax"] = m["tax_no"] or ""
                st.session_state[f"med_p{i}_addr"] = m["address"] or ""
                st.session_state[f"med_p{i}_att"] = m["attorney"] or ""
                st.session_state[f"med_p{i}_bar"] = m["attorney_bar"] or ""
                st.rerun()

    ptype = st.selectbox("Tür", ["Gerçek Kişi","Şirket"], key=f"med_p{i}_type")
    name = st.text_input("Ad Soyad / Unvan", key=f"med_p{i}_name")
    c1,c2 = st.columns(2)
    with c1:
        identity = st.text_input("T.C. Kimlik No", key=f"med_p{i}_id", disabled=(ptype=="Şirket"))
    with c2:
        tax = st.text_input("Vergi No", key=f"med_p{i}_tax", disabled=(ptype=="Gerçek Kişi"))
    address = st.text_area("Adres", key=f"med_p{i}_addr", height=80)
    attorney = st.text_input("Vekili", key=f"med_p{i}_att", placeholder="Av. ...")
    bar = st.text_input("Baro / Sicil", key=f"med_p{i}_bar", placeholder="İstanbul Baro 12345")

    p = {
        "party_type":ptype, "name":name, "identity_no":identity,
        "tax_no":tax, "address":address, "attorney":attorney, "attorney_bar":bar
    }
    if st.button("Tarafı kaydet", key=f"med_save_{i}"):
        save_party(p)
        st.success("Taraf kaydedildi.")
    return p


def render_mediation():
    st.header("Arabuluculuk")
    st.caption("Prepare • Generate • Word + PDF")

    mediator_choice = st.selectbox("Arabulucu", list(MEDIATORS.keys()))
    mediator = MEDIATORS[mediator_choice]
    c1,c2 = st.columns([1,3])
    c1.metric("Sicil", mediator["registry"])
    c2.write(f"**Adres:** {mediator['address']}")

    st.subheader("Taraflar")
    party_count = st.slider("Taraf sayısı", 1, 5, 2)
    parties = []
    for i in range(1, party_count+1):
        with st.expander(f"Taraf {i}", expanded=(i <= 2)):
            parties.append(_party_editor(i))

    st.divider()
    st.subheader("Dosya Bilgileri")
    c1,c2 = st.columns(2)
    file_no = c1.text_input("Dosya No")
    kind = c2.selectbox(
        "Uyuşmazlık Türü",
        ["İşçilik Alacağı","Kira Tespit","Kiralananın Tahliyesi","Kira Tespit + Kiralananın Tahliyesi"]
    )
    d1,d2,d3 = st.columns(3)
    opening_date = d1.date_input("Dosya açılış tarihi", value=None, format="DD.MM.YYYY")
    agreement_date = d2.date_input("Anlaşma tarihi", value=None, format="DD.MM.YYYY")
    final_date = d3.date_input("Son tutanak tarihi", value=None, format="DD.MM.YYYY")

    st.divider()
    st.subheader("İşten Çıkış / Ayrılış Bildirgesi")
    exit_doc = st.file_uploader("Belgeyi yükle; taraf bilgilerini taslağa önerebilir", type=["pdf","docx"], key="med_exit")
    if exit_doc:
        txt = extract_upload_text(exit_doc)
        if txt.strip():
            if st.button("Belgeden taraf bilgilerini çıkar"):
                try:
                    parsed = ai_extract_exit_notice(txt)
                    st.session_state["med_exit_parsed"] = parsed
                except Exception as e:
                    st.error(f"Belge okunamadı: {e}")
            parsed = st.session_state.get("med_exit_parsed")
            if parsed:
                st.json(parsed)
                a,b = st.columns(2)
                if a.button("İşçiyi Taraf 1'e aktar"):
                    emp = parsed.get("employee",{})
                    st.session_state["med_p1_type"]="Gerçek Kişi"
                    st.session_state["med_p1_name"]=emp.get("name","")
                    st.session_state["med_p1_id"]=emp.get("identity_no","")
                    st.session_state["med_p1_addr"]=emp.get("address","")
                    st.rerun()
                if b.button("İşvereni Taraf 2'ye aktar"):
                    er = parsed.get("employer",{})
                    st.session_state["med_p2_type"]="Şirket"
                    st.session_state["med_p2_name"]=er.get("name","")
                    st.session_state["med_p2_tax"]=er.get("tax_no","")
                    st.session_state["med_p2_addr"]=er.get("address","")
                    st.rerun()
        else:
            st.warning("Bu belgeden metin çıkarılamadı. V1'de metin tabanlı PDF/DOCX destekleniyor.")

    st.divider()
    st.subheader("Anlaşma Taslak Alanları")
    details = {}
    if kind == "İşçilik Alacağı":
        c1,c2 = st.columns(2)
        details["termination_date"] = c1.date_input("İş sözleşmesi sona erme tarihi", value=None, format="DD.MM.YYYY")
        details["receivable"] = c2.text_input("Anlaşmaya konu alacak", placeholder="Örn. Kıdem tazminatı")
        c3,c4,c5 = st.columns(3)
        details["amount"] = c3.text_input("Tutar", placeholder="Boş bırakılabilir")
        details["netgross"] = c4.selectbox("Net / Brüt", ["","NET","BRÜT"])
        details["payment_date"] = c5.date_input("Ödeme tarihi", value=None, format="DD.MM.YYYY")
        details["payment_method"] = st.text_input("Ödeme yöntemi", placeholder="Örn. banka hesabına / nakden")
    else:
        details["property_address"] = st.text_area("Taşınmaz adresi", height=70)
        c1,c2,c3 = st.columns(3)
        details["lease_start"] = c1.date_input("Kira sözleşmesi başlangıcı", value=None, format="DD.MM.YYYY")
        details["lease_end"] = c2.date_input("Kira sözleşmesi sona erme tarihi", value=None, format="DD.MM.YYYY")
        details["eviction_date"] = c3.date_input("Tahliye / teslim tarihi", value=None, format="DD.MM.YYYY")
        st.write("**Kira dönemleri**")
        rent_periods=[]
        for i in range(1,4):
            r1,r2,r3 = st.columns([2,1,2])
            period = r1.text_input(f"Dönem {i}", key=f"rent_period_{i}", placeholder="Örn. 2026 Ağustos-Aralık")
            amount = r2.text_input(f"Aylık tutar {i}", key=f"rent_amount_{i}")
            rule = r3.text_input(f"Ödeme şekli {i}", key=f"rent_rule_{i}", placeholder="her ay peşin ve eksiksiz")
            rent_periods.append({"period":period,"amount":amount,"payment_rule":rule})
        details["rent_periods"]=rent_periods

    st.write("**Arabuluculuk ücreti**")
    f1,f2,f3 = st.columns(3)
    details["mediation_fee"] = f1.text_input("Arabuluculuk ücreti", placeholder="Boş bırakılabilir")
    details["fee_payer"] = f2.text_input("Ücreti ödeyecek taraf", placeholder="Örn. TARAF 2 ...")
    details["fee_payment_date"] = f3.date_input("Ücret ödeme tarihi", value=None, format="DD.MM.YYYY")
    details["fee_account"] = st.text_input("Ödeme hesabı / IBAN", placeholder="Boş bırakılabilir")
    details["fee_plus_vat"] = st.checkbox("Arabuluculuk ücretine + KDV ekle")

    data = {
        "mediator":mediator, "parties":parties, "file_no":file_no,
        "opening_date":opening_date, "agreement_date":agreement_date, "final_date":final_date,
        "dispute_kind":kind, "details":details
    }

    if st.button("🧾 3 Belge Taslağını Oluştur", type="primary", use_container_width=True):
        st.session_state["med_info_text"] = build_information_text(data)
        st.session_state["med_agreement_text"] = build_agreement_text(data)
        st.session_state["med_final_text"] = build_final_text(data)

    if st.session_state.get("med_info_text"):
        st.divider()
        st.header("Taslakları Düzenle")
        st.caption("Rakam ve tarihler boş bırakılabilir. İndirmeden önce metnin tamamını burada değiştirebilirsin.")

        t1,t2,t3 = st.tabs(["Bilgilendirme ve Belirleme","Anlaşma Belgesi","Son Tutanak"])
        with t1:
            info_text = st.text_area("Bilgilendirme ve Belirleme Tutanağı", st.session_state["med_info_text"], height=600)
            st.session_state["med_info_text"] = info_text
        with t2:
            agreement_text = st.text_area("Anlaşma Belgesi", st.session_state["med_agreement_text"], height=650)
            st.session_state["med_agreement_text"] = agreement_text
        with t3:
            final_text = st.text_area("Son Tutanak", st.session_state["med_final_text"], height=600)
            st.session_state["med_final_text"] = final_text

        docs = [
            ("Bilgilendirme ve Belirleme Tutanağı", st.session_state["med_info_text"], "bilgilendirme_belirleme"),
            ("Anlaşma Belgesi", st.session_state["med_agreement_text"], "anlasma_belgesi"),
            ("Son Tutanak", st.session_state["med_final_text"], "son_tutanak"),
        ]
        for title, text, stem in docs:
            st.subheader(title)
            c1,c2 = st.columns(2)
            c1.download_button(
                f"⬇️ {title} - Word",
                data=text_to_docx(text),
                file_name=f"{file_no or 'taslak'} - {stem}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key=f"docx_{stem}"
            )
            c2.download_button(
                f"⬇️ {title} - PDF",
                data=text_to_pdf(text),
                file_name=f"{file_no or 'taslak'} - {stem}.pdf",
                mime="application/pdf",
                use_container_width=True,
                key=f"pdf_{stem}"
            )

        st.warning("V1 taraf hafızası prototip SQLite kullanır. Streamlit yeniden deploy edildiğinde kayıtlar sıfırlanabilir; kalıcı ve şifreli taraf veritabanını güvenlik katmanıyla birlikte kuracağız.")
