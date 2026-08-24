from io import BytesIO
import re
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX

AUTHOR = "Av. Onur Güneş"

def _norm(s):
    return re.sub(r"\s+"," ",(s or "")).strip()

def _all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def _ref_match(ptext, ref):
    if not ref: return False
    p=_norm(ptext); r=_norm(ref)
    # Supports "5.01", "5.01 / 5.02", "Madde 5.01" etc.
    refs=re.findall(r"\d+(?:\.\d+)+",r)
    return any(re.search(rf"(^|\b){re.escape(x)}(?:[\.\s:\-]|$)",p,re.I) for x in refs)

def _highlight_first_visible_token(paragraph):
    """Highlight only the first visible token/run in yellow."""
    for run in paragraph.runs:
        if run.text and run.text.strip():
            txt=run.text
            m=re.search(r"\S+",txt)
            if not m: continue
            # Split run so only first token is highlighted.
            before=txt[:m.start()]
            token=txt[m.start():m.end()]
            after=txt[m.end():]
            run.text=before
            new=paragraph.add_run(token)
            new.font.highlight_color=WD_COLOR_INDEX.YELLOW
            if after:
                paragraph.add_run(after)
            # Move inserted runs directly after original run.
            p=paragraph._p
            elems=list(p)
            # python-docx append puts at end; reconstructing exact visual order is safer:
            added=[r._r for r in paragraph.runs if r.text in (token,after) and r._r is not run._r]
            # If ordering cannot be safely inferred, fallback below is harmless.
            return True
    return False

def _highlight_first_word_simple(paragraph):
    # Robust fallback: highlight the first run containing visible text.
    for run in paragraph.runs:
        if run.text and run.text.strip():
            run.font.highlight_color=WD_COLOR_INDEX.YELLOW
            return True
    return False

def _tracked_replace_whole_paragraph(paragraph, new_text, change_id):
    """Tracked delete+insert while preserving the paragraph's existing run formatting."""
    old_text=paragraph.text

    # Capture formatting from the first visible original run.
    template_rpr=None
    for run in paragraph.runs:
        if run.text and run.text.strip():
            rpr=run._r.find(qn("w:rPr"))
            if rpr is not None:
                template_rpr=deepcopy(rpr)
            break

    # Remove content but preserve paragraph properties.
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)

    dele=OxmlElement("w:del")
    dele.set(qn("w:id"),str(change_id)); change_id+=1
    dele.set(qn("w:author"),AUTHOR)
    dr=OxmlElement("w:r")
    if template_rpr is not None:
        dr.append(deepcopy(template_rpr))
    dt=OxmlElement("w:delText")
    dt.set(qn("xml:space"),"preserve")
    dt.text=old_text
    dr.append(dt); dele.append(dr); paragraph._p.append(dele)

    ins=OxmlElement("w:ins")
    ins.set(qn("w:id"),str(change_id)); change_id+=1
    ins.set(qn("w:author"),AUTHOR)
    ir=OxmlElement("w:r")
    if template_rpr is not None:
        ir.append(deepcopy(template_rpr))
    it=OxmlElement("w:t")
    it.set(qn("xml:space"),"preserve")
    it.text=new_text
    ir.append(it); ins.append(ir); paragraph._p.append(ins)
    return change_id

def _tracked_replace_excerpt(paragraph, excerpt, suggestion, change_id):
    """
    Try a micro replacement using normalized exact excerpt.
    If excerpt is not literally present, caller may fall back to paragraph-level replacement.
    """
    ptext=paragraph.text
    if excerpt and excerpt in ptext:
        new_text=ptext.replace(excerpt,suggestion,1)
        return _tracked_replace_whole_paragraph(paragraph,new_text,change_id), True
    return change_id, False

def _find_target(paragraphs, rev):
    excerpt=_norm(rev.get("current_excerpt",""))
    ref=rev.get("reference","")
    # First: exact/normalized excerpt containment.
    if excerpt:
        for p in paragraphs:
            if excerpt in _norm(p.text):
                return p
    # Second: real clause/reference.
    for p in paragraphs:
        if _ref_match(p.text,ref):
            return p
    return None

def create_revised_word(raw, revisions):
    doc=Document(BytesIO(raw))
    paragraphs=list(_all_paragraphs(doc))
    change_id=1
    applied=0
    attention=0

    for rev in revisions:
        target=_find_target(paragraphs,rev)
        if target is None:
            attention+=1
            continue

        suggestion=_norm(rev.get("new_text","") or rev.get("suggested_revision",""))
        excerpt=_norm(rev.get("old_text","") or rev.get("current_excerpt",""))
        rtype=(rev.get("revision_type") or "").upper()

        # Every OLI attention point gets ONE yellow marker only.
        _highlight_first_word_simple(target)

        if rtype in ("EK HÜKÜM","SİLME") or not suggestion:
            attention+=1
            continue

        # Prefer micro/excerpt replacement if literal text is available.
        cid,ok=_tracked_replace_excerpt(target,excerpt,suggestion,change_id)
        if ok:
            change_id=cid; applied+=1
            continue

        # If model supplied a usable complete revision but excerpt mapping was imperfect,
        # apply it to the referenced paragraph as tracked replacement.
        if suggestion:
            change_id=_tracked_replace_whole_paragraph(target,suggestion,change_id)
            applied+=1
        else:
            attention+=1

    bio=BytesIO()
    doc.save(bio)
    return bio.getvalue(), {"applied":applied,"attention":attention}
